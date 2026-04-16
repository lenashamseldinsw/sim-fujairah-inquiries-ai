"""Automatically detect report structure from Word documents."""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, List, Any, Tuple, Optional
import re


class ReportStructureDetector:
    """
    Automatically detects sections, tables, and charts from Word documents.

    This makes the system adaptable to different report formats without hardcoding.
    """

    # Common heading patterns (Arabic and English)
    HEADING_PATTERNS = [
        r'(التحليل|تحليل|Analysis)',
        r'(الملخص|ملخص|Summary)',
        r'(الخلاصة|خلاصة|Conclusion)',
        r'(التوصيات|توصيات|Recommendations)',
        r'(المقدمة|مقدمة|Introduction)',
        r'(حالات الاستخدام|Use Cases)',
        r'(خارطة الطريق|Roadmap)',
    ]

    def __init__(self, doc_path: str):
        """Initialize detector with document path."""
        self.doc_path = doc_path
        self.doc = Document(doc_path)

    def detect_structure(self) -> Dict[str, Any]:
        """
        Detect complete report structure with hierarchical sections and subsections.

        Returns:
            Dictionary with main sections, subsections, tables, and metadata
        """
        # First detect all headings (main + sub)
        print("\n📄 Detecting document structure...")
        all_headings = self._detect_all_headings()
        print(f"  ✓ Found {len(all_headings)} headings")

        # Organize into hierarchical structure
        main_sections = self._organize_hierarchy(all_headings)
        print(f"  ✓ Organized into {len(main_sections)} main sections")
        for sec in main_sections:
            subsec_count = len(sec.get('subsections', []))
            content_len = len(sec.get('content', ''))
            print(f"    - {sec['title_ar']} ({subsec_count} subsections, {content_len} chars content)")

        # Detect tables
        print("\n📊 Detecting tables...")
        tables = self._detect_tables()
        print(f"  ✓ Found {len(tables)} tables total")

        # Detect charts
        charts = self._detect_chart_positions()
        print(f"  ✓ Found {len(charts)} charts")

        # Assign tables and charts to sections/subsections
        print("\n🔗 Assigning elements to sections...")
        self._assign_elements_to_hierarchy(main_sections, tables, charts)

        # Debug: Report unassigned tables
        unassigned_tables = [t for t in tables if t.get('assigned_section') is None]
        if unassigned_tables:
            print(f"  ⚠️  {len(unassigned_tables)} tables could not be assigned!")
        else:
            print(f"  ✓ All tables assigned successfully")

        return {
            'sections': main_sections,
            'tables': tables,
            'charts': charts,
            'metadata': self._extract_metadata()
        }

    def _detect_all_headings(self) -> List[Dict[str, Any]]:
        """
        Detect all headings (main sections and subsections).

        Single-pass through document body to find headings with correct positions.
        This ensures position consistency with tables and charts.

        Returns:
            List of all headings with their metadata
        """
        headings = []
        element_position = 0
        heading_count = 0

        # Single pass through document body
        # This gives us authoritative element positions
        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                # Get paragraph text
                text_elems = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                para_text = ''.join([t.text for t in text_elems if t.text]).strip()

                # Check if this paragraph is a heading
                heading_type = self._classify_heading_from_text(para_text, element)

                if heading_type:
                    heading = {
                        'id': self._generate_section_id(para_text, heading_count),
                        'title_ar': para_text,
                        'title_en': self._translate_section_title(para_text),
                        'position': element_position,  # Authoritative position in body
                        'type': heading_type,
                        'level': self._get_heading_level_from_element(element),
                        'tables': [],
                        'charts': [],
                        'subsections': [],
                        'content': ''
                    }
                    headings.append(heading)
                    heading_count += 1

            # Count every element for consistent positioning
            element_position += 1

        return headings

    def _classify_heading_from_text(self, text: str, element) -> str:
        """
        Classify heading from text and element style.

        Detects main sections and subsections while avoiding false positives from
        regular content paragraphs.
        """
        if not text:
            return None

        # Bullet points and list items should NOT be classified as headings
        # These are actual content that should appear in the section
        if text.startswith('←') or text.startswith('-') or text.startswith('•') or text.startswith('▪'):
            return None

        # Also check for complaint type patterns (from section 2.2)
        # These look like: "الشكوى العاجلة\الحرجة: ذات الإجراءات الواضحة — إغلاق خلال 24 ساعة"
        if (text.startswith('الشكوى ') and '—' in text and
            ('خلال' in text and ('ساعة' in text or 'أيام' in text or 'يوم' in text))):
            return None

        # Question paragraphs (starting with لماذا, هل, ما, etc.) are content, not headings
        if any(text.startswith(q) for q in ['لماذا', 'هل', 'ما ', 'من ', 'أين ', 'متى ', 'كيف ']):
            return None

        # Check for bold ending FIRST (before length checks that would discard long paragraphs)
        # This allows us to detect hidden headings at the end of long paragraphs
        bold_ending = self._extract_bold_ending(element)
        if bold_ending:
            # Debug: log what we found
            if 'حل الشكاوى' in bold_ending or 'الاكتشاف' in bold_ending:
                print(f"  🔍 DEBUG: Found potential heading: '{bold_ending}' (len={len(bold_ending)}, ends_with_period={bold_ending.endswith('.')})")

            if len(bold_ending) < 100:
                # Bold ending is short - check if it looks like a heading (not a sentence)
                if not bold_ending.endswith('.'):
                    if 'حل الشكاوى' in bold_ending or 'الاكتشاف' in bold_ending:
                        print(f"  ✓ DEBUG: Classified as heading!")
                    return 'sub'

        # Now check full text (reject if too long for standard headings)
        # Question paragraphs and explanation text are usually longer
        if len(text) > 300:
            return None

        # Check font size (skip TOC entries)
        for run in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            if rPr is not None:
                sz = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
                if sz is not None:
                    try:
                        font_size = int(sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')) / 2
                        if font_size < 11:
                            return None
                    except:
                        pass

        # Check style
        pPr = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            pStyle = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if pStyle is not None:
                style_val = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                # Handle both "Heading 1"/"Heading1" and "Heading 2"/"Heading2" formats
                if 'Heading' in style_val:
                    # Remove spaces for comparison (Heading 1 vs Heading1)
                    normalized = style_val.replace(' ', '')
                    if 'Heading1' in normalized:
                        return 'main'
                    else:
                        return 'sub'

        # Pattern matching
        import re
        # Arabic ordinals at start with colon (main sections)
        # Handle variations: with/without diacritics, with/without space before colon
        if re.search(r'^(أولاً|أولا|ثانياً|ثانيا|ثالثاً|ثالثا|رابعاً|رابعا|خامساً|خامسا|سادساً|سادسا|سابعاً|سابعا|ثامناً|ثامنا|تاسعاً|تاسعا|عاشراً|عاشرا)\s*:', text):
            return 'main'

        # Also check for ordinals WITHOUT colon (style variation)
        if re.search(r'^(أولاً|أولا|ثانياً|ثانيا|ثالثاً|ثالثا|رابعاً|رابعا|خامساً|خامسا|سادساً|سادسا|سابعاً|سابعا|ثامناً|ثامنا|تاسعاً|تاسعا|عاشراً|عاشرا)\s', text):
            return 'main'

        # Numbered subsections: "X.Y " format (must have space after)
        if re.search(r'^\d+\.\d+\s', text):
            return 'sub'

        # Also catch "X.Y:" format (colon instead of space)
        if re.search(r'^\d+\.\d+:', text):
            return 'sub'

        # Catch "X. " format for numbered sections (like "1. ", "2. ", etc. for first-level items)
        # But be careful not to catch paragraphs that happen to start with numbers
        # Only if the number is 1-10 (typical for main section numbers) and followed by space/colon
        if re.search(r'^[1-9]\.\s+[A-Z\u0600-\u06FF]', text) or re.search(r'^[1-9]:\s+[A-Z\u0600-\u06FF]', text):
            return 'sub'

        # Check for Arabic section marker patterns (e.g., "الجزء الخامس:", "الفصل 5:")
        if re.search(r'^(الجزء|الفصل|الباب|القسم)\s+(أول|ثاني|ثالث|رابع|خامس|سادس|سابع|ثامن|تاسع|عاشر|\d+)[:\s]', text):
            return 'main'

        return None

    def _extract_bold_ending(self, para_element) -> str:
        """
        Extract bold text from the end of a paragraph.

        Concatenates all runs from the last bold run onwards.
        Returns empty string if no bold runs found.
        """
        runs = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        if not runs:
            return ''

        # Find the index of the last bold run
        last_bold_idx = -1
        for i in range(len(runs) - 1, -1, -1):
            rPr = runs[i].find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            if rPr is not None:
                b = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                if b is not None:
                    last_bold_idx = i
                    break

        if last_bold_idx < 0:
            return ''  # No bold runs

        # Extract text from last bold run onwards
        bold_text = []
        for i in range(last_bold_idx, len(runs)):
            text_elems = runs[i].findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            text = ''.join([t.text for t in text_elems if t.text])
            if text:
                bold_text.append(text)

        result = ''.join(bold_text).strip()
        if 'حل الشكاوى' in result:
            print(f"  🔍 DEBUG _extract_bold_ending: Found '{result}' (from run index {last_bold_idx})")
        return result

    def _extract_non_bold_prefix(self, para_element) -> str:
        """
        Extract all non-bold text from a paragraph, stopping at first bold text.

        Used to separate regular content from bold headings that might appear at the end.

        Returns the concatenated text of all runs before the first bold run.
        """
        runs = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        non_bold_text = []

        for run in runs:
            rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            is_bold = False

            if rPr is not None:
                b = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                if b is not None:
                    is_bold = True

            if is_bold:
                # Stop at first bold run
                break

            # Extract text from this non-bold run
            text_elems = run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            text = ''.join([t.text for t in text_elems if t.text])
            if text:
                non_bold_text.append(text)

        return ''.join(non_bold_text)

    def _is_text_bold(self, element) -> bool:
        """
        Check if text in element is bold.

        Returns True if most runs are bold, or if the last run(s) are bold with short text.
        """
        runs = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        if not runs:
            return False

        bold_count = 0
        for run in runs:
            rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            if rPr is not None:
                b = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                if b is not None:
                    bold_count += 1

        # Consider it bold if most runs are bold
        if bold_count > len(runs) / 2:
            return True

        # Also check if the LAST run(s) are bold (may be a bold heading at end of paragraph)
        if runs:
            last_run = runs[-1]
            rPr = last_run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            if rPr is not None:
                b = rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                if b is not None:
                    return True

        return False

    def _get_heading_level_from_element(self, element) -> int:
        """Get heading level from element."""
        pPr = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            pStyle = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if pStyle is not None:
                style_val = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                import re
                match = re.search(r'Heading (\d+)', style_val)
                if match:
                    return int(match.group(1))
        return 2  # Default

    def _organize_hierarchy(self, headings: List[Dict]) -> List[Dict]:
        """
        Organize flat list of headings into hierarchical structure.

        Main sections contain their subsections, matched by section number.
        Removes duplicate main sections (e.g., from table of contents).
        Automatically skips cover page and table of contents before first main section.

        Returns:
            List of main sections with nested subsections
        """
        if not headings:
            return []

        # Find the position of the first "real" main section (skip TOC)
        # Real main sections start with "أولاً:" or similar Arabic ordinals
        first_real_main_pos = None
        for i, heading in enumerate(headings):
            if heading['type'] == 'main':
                # Check if it matches the main section pattern (Arabic ordinal with colon)
                # Handle variations: with/without diacritics, with/without space before colon
                import re
                if re.search(r'^(أولاً|أولا|ثانياً|ثانيا|ثالثاً|ثالثا|رابعاً|رابعا|خامساً|خامسا|سادساً|سادسا|سابعاً|سابعا|ثامناً|ثامنا|تاسعاً|تاسعا|عاشراً|عاشرا)\s*:', heading['title_ar']):
                    first_real_main_pos = i
                    break

        # If no real main section found, use first heading
        if first_real_main_pos is None:
            first_real_main_pos = 0

        # Filter out cover page / TOC - everything before first real main section
        filtered_headings = headings[first_real_main_pos:]

        # Also store the position where main content starts (for skipping cover/TOC in content extraction)
        if filtered_headings:
            self._main_content_start_pos = filtered_headings[0]['position']
            if first_real_main_pos > 0:
                print(f"\n  🔖 Skipping cover/TOC: {first_real_main_pos} headings before first main section")
        else:
            self._main_content_start_pos = 0

        # First, collect all main sections and subsections separately
        main_sections_by_title = {}
        main_sections_order = []
        subsections = []

        for heading in filtered_headings:
            if heading['type'] == 'main':
                title_key = heading['title_ar'].strip()

                # Only keep first occurrence (skip table of contents duplicates)
                if title_key not in main_sections_by_title:
                    heading['subsections'] = []
                    main_sections_by_title[title_key] = heading
                    main_sections_order.append(heading)

            elif heading['type'] == 'sub':
                subsections.append(heading)

        # Now assign subsections to their parent main section by matching section numbers
        # E.g., "2.1" belongs to "ثانياً", "3.2" belongs to "ثالثاً", etc.
        # Include variations with and without diacritics
        arabic_ordinals = {
            'أولاً': 1, 'أولا': 1,
            'ثانياً': 2, 'ثانيا': 2,
            'ثالثاً': 3, 'ثالثا': 3,
            'رابعاً': 4, 'رابعا': 4,
            'خامساً': 5, 'خامسا': 5,
            'سادساً': 6, 'سادسا': 6,
            'سابعاً': 7, 'سابعا': 7,
            'ثامناً': 8, 'ثامنا': 8,
            'تاسعاً': 9, 'تاسعا': 9,
            'عاشراً': 10, 'عاشرا': 10,
        }

        # Create mapping from section number to main section
        main_sections_by_number = {}
        for main_sec in main_sections_order:
            for ordinal, num in arabic_ordinals.items():
                if ordinal in main_sec['title_ar']:
                    main_sections_by_number[num] = main_sec
                    break

        # Assign subsections based on their number prefix (e.g., "2.1" -> section 2)
        # Or by position if no number prefix
        unnumbered_subsecs = []

        for subsec in subsections:
            # Extract section number from subsection title (e.g., "2.1" -> 2)
            # Try multiple patterns to be more robust
            match = re.match(r'^(\d+)\.', subsec['title_ar'])
            if match:
                section_num = int(match.group(1))
                parent_section = main_sections_by_number.get(section_num)

                if parent_section:
                    parent_section['subsections'].append(subsec)
                    print(f"    → Assigned subsection: {subsec['title_ar']}")
                else:
                    print(f"    ⚠️  Subsection {subsec['title_ar']} has no parent section {section_num}")
            else:
                # No number prefix - assign by position (closest preceding main section)
                unnumbered_subsecs.append(subsec)

        # Assign unnumbered subsections to their parent by position proximity
        for subsec in unnumbered_subsecs:
            # Find closest preceding main section by position
            closest_parent = None
            closest_distance = float('inf')

            for main_sec in main_sections_order:
                if main_sec['position'] < subsec['position']:
                    distance = subsec['position'] - main_sec['position']
                    if distance < closest_distance:
                        closest_distance = distance
                        closest_parent = main_sec

            if closest_parent:
                closest_parent['subsections'].append(subsec)
                print(f"    → Assigned unnumbered subsection to {closest_parent['title_ar']}: {subsec['title_ar']}")
            else:
                print(f"    ⚠️  Unnumbered subsection has no parent: {subsec['title_ar']}")

        # Remove duplicate main sections - keep the one with subsections
        # Also filter out non-ordinal sections (like document title)
        final_sections = []
        seen_numbers = set()  # Track by Arabic ordinal number

        # Include both diacritic variants when checking for ordinals
        arabic_ordinals_list = ['أولاً', 'أولا', 'ثانياً', 'ثانيا', 'ثالثاً', 'ثالثا',
                                'رابعاً', 'رابعا', 'خامساً', 'خامسا',
                                'سادساً', 'سادسا', 'سابعاً', 'سابعا',
                                'ثامناً', 'ثامنا', 'تاسعاً', 'تاسعا', 'عاشراً', 'عاشرا']

        for main_sec in main_sections_order:
            # Check if this section has an Arabic ordinal
            ordinal_found = None
            for ordinal in arabic_ordinals_list:
                if ordinal in main_sec['title_ar']:
                    ordinal_found = ordinal
                    break

            # Only keep sections with Arabic ordinals (filters out title page, etc.)
            if ordinal_found:
                if ordinal_found not in seen_numbers:
                    final_sections.append(main_sec)
                    seen_numbers.add(ordinal_found)
                else:
                    # This ordinal already exists - merge if this one has more subsections
                    existing_sec = next(s for s in final_sections if ordinal_found in s['title_ar'])
                    if len(main_sec['subsections']) > len(existing_sec['subsections']):
                        # This duplicate has more subsections, replace
                        idx = final_sections.index(existing_sec)
                        final_sections[idx] = main_sec
                    elif main_sec['subsections']:
                        # Both have subsections, merge them
                        existing_sec['subsections'].extend(main_sec['subsections'])

        # Extract content for each section and subsection
        self._extract_content_for_hierarchy(final_sections)

        return final_sections

    def _extract_content_for_hierarchy(self, main_sections: List[Dict]) -> None:
        """
        Extract content for main sections and subsections using dynamic proximity-based assignment.

        Each paragraph is assigned to the closest preceding section (main or subsection).
        This handles edge cases and orphaned content automatically without position ranges.
        """
        # Create a mapping of element body positions to paragraphs
        element_pos_to_para = {}
        all_para_positions = []  # All paragraph positions in order
        element_position = 0

        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                # Store paragraph element and track its position
                element_pos_to_para[element_position] = element
                all_para_positions.append(element_position)
            element_position += 1

        # Build a position map: position -> (section, is_main)
        # This maps each section's position to the section object
        position_to_section = {}
        section_positions = []  # Sorted list of section positions

        for main_sec in main_sections:
            pos = main_sec['position']
            position_to_section[pos] = (main_sec, True)  # is_main=True
            section_positions.append(pos)

            for subsec in main_sec['subsections']:
                pos = subsec['position']
                position_to_section[pos] = (subsec, False)  # is_main=False
                section_positions.append(pos)

        # Sort section positions
        section_positions.sort()

        # Initialize content for all sections
        for main_sec in main_sections:
            main_sec['content'] = ''
            for subsec in main_sec['subsections']:
                subsec['content'] = ''

        # Get main content start position (for skipping cover/TOC)
        main_content_start = getattr(self, '_main_content_start_pos', 0)

        # Process each paragraph and assign to closest preceding section
        for para_pos in all_para_positions:
            # Skip paragraphs before main content (cover page, TOC)
            if para_pos < main_content_start:
                continue

            if para_pos not in element_pos_to_para:
                continue

            para_element = element_pos_to_para[para_pos]

            # Get text from paragraph
            text_elems = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            full_text = ''.join([t.text for t in text_elems if t.text]).strip()

            if not full_text:
                continue  # Skip empty paragraphs

            # Check if this paragraph is itself a heading (shouldn't be assigned as content)
            is_heading = self._classify_heading_from_text(full_text, para_element)
            if is_heading:
                continue  # Skip headings

            # Find the closest preceding section
            closest_section_pos = None
            for sec_pos in section_positions:
                if sec_pos < para_pos:
                    closest_section_pos = sec_pos
                else:
                    break  # Stop at first section at or after this paragraph

            if closest_section_pos is None:
                # No section before this paragraph, assign to first section
                if main_sections:
                    closest_section_pos = main_sections[0]['position']
                else:
                    continue  # No sections at all, skip

            # Get the section object
            section, is_main = position_to_section[closest_section_pos]

            # Process paragraph content
            para_content = self._process_paragraph_content(full_text, para_element)

            if para_content:
                # Append to section's content
                if section['content']:
                    section['content'] += '\n' + para_content
                else:
                    section['content'] = para_content

    def _process_paragraph_content(self, full_text: str, para_element) -> str:
        """
        Process paragraph content, handling bullet points and bold text appropriately.

        Returns the content to include in section (may strip bold endings that look like headings).
        """
        # Bullet points should always include full text
        if full_text.startswith('←') or full_text.startswith('-') or full_text.startswith('•'):
            return full_text

        # For regular paragraphs, check if they end with bold text that might be a heading
        text_without_bold_ending = self._extract_non_bold_prefix(para_element)

        if text_without_bold_ending and len(text_without_bold_ending) < len(full_text):
            # Paragraph contains bold ending
            bold_ending = self._extract_bold_ending(para_element)

            # If bold ending is short and doesn't end with period, it might be a heading for next section
            if bold_ending and len(bold_ending) < 100 and not bold_ending.endswith('.'):
                # Only include non-bold part
                return text_without_bold_ending.strip() if text_without_bold_ending.strip() else ""
            else:
                # Bold ending is not a heading, include full text
                return full_text
        else:
            # No bold ending, include full text
            return full_text

    def _detect_tables(self) -> List[Dict[str, Any]]:
        """
        Detect all tables with their position and captions in document.

        Dynamically detects table captions by looking at text immediately preceding each table.
        Captions are descriptive text (not section headings) that appear right before a table.
        Automatically skips tables that appear before main content (cover/TOC pages).

        Returns:
            List of table metadata with position, data, and optional captions
        """
        tables_info = []

        for table_idx, table in enumerate(self.doc.tables):
            # Find table position by looking at surrounding elements
            table_position = self._find_table_position(table, table_idx)

            # Skip tables before main content starts (cover page, TOC)
            if hasattr(self, '_main_content_start_pos') and table_position < self._main_content_start_pos:
                print(f"  📋 Table {table_idx} at position {table_position} SKIPPED (before main content)")
                continue

            # Extract table data
            table_data = self._extract_table_data(table)

            # Dynamically detect caption by looking at preceding elements
            caption = self._extract_table_caption(table_idx, table_position)

            # Debug: Log table detection
            rows_count = len(table.rows)
            cols_count = len(table.rows[0].cells) if table.rows else 0
            caption_marker = f" [{caption[:40]}...]" if caption else ""
            print(f"  📋 Table {table_idx} detected at position {table_position} ({rows_count}x{cols_count}){caption_marker}")

            tables_info.append({
                'index': table_idx,
                'position': table_position,
                'data': table_data,
                'caption': caption,  # Dynamically detected caption
                'assigned_section': None  # Will be assigned later
            })

        return tables_info

    def _identify_caption_positions(self, element_pos_to_para: Dict[int, Any], table_positions: set) -> set:
        """
        Identify which paragraph positions are table captions.

        A position is a caption if:
        - It's between a heading and a table
        - It's not a heading itself
        - It's reasonably short (< 250 chars)
        - It's not a bullet point or list

        Captions are typically the paragraph(s) immediately before a table that describe it.

        Returns:
            Set of element positions that are table captions
        """
        caption_positions = set()

        for table_pos in sorted(table_positions):
            # Look backwards from table to find all text before it
            # Stop when we hit a heading (section boundary)
            candidate_captions = []
            first_heading_pos = None

            # Check positions before the table
            for check_pos in range(table_pos - 1, max(-1, table_pos - 10), -1):
                if check_pos not in element_pos_to_para:
                    continue

                para_element = element_pos_to_para[check_pos]

                # Get text from this paragraph
                text_elems = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                text = ''.join([t.text for t in text_elems if t.text]).strip()

                if not text:
                    continue  # Skip empty paragraphs

                # Check if it's a heading
                is_heading = self._classify_heading_from_text(text, para_element)
                if is_heading:
                    # Found a heading - this is a section boundary
                    # Everything we found between this heading and the table are potential captions
                    first_heading_pos = check_pos
                    break

                # Skip if too long (likely regular paragraph content, not a caption)
                if len(text) > 250:
                    continue

                # Skip bullet points or lists
                if text.startswith('←') or text.startswith('-') or text.startswith('•'):
                    continue

                # This is a candidate caption
                candidate_captions.append((check_pos, text))

            # Add the closest paragraph to the table as a caption (usually the actual caption)
            if candidate_captions:
                # The last one in the list is closest to the table
                caption_pos, _ = candidate_captions[0]
                caption_positions.add(caption_pos)

        return caption_positions

    def _extract_table_caption(self, table_idx: int, table_position: int) -> Optional[str]:
        """
        Dynamically extract caption for a table by examining preceding paragraphs.

        Captions are:
        - Non-empty text
        - Not section headings
        - Immediately before the table
        - Typically descriptive or indicative of table content

        Returns:
            Caption text if found, None otherwise
        """
        # Look back from table position to find preceding text
        preceding_text_paras = []

        # Scan backwards through body elements before the table
        current_elem_pos = 0
        for element in self.doc.element.body:
            if current_elem_pos >= table_position:
                break

            if element.tag.endswith('p'):
                preceding_text_paras.append((current_elem_pos, element))

            current_elem_pos += 1

        # Examine the paragraphs immediately before the table
        # Typically a caption would be 1-2 paragraphs before
        candidates = []
        for elem_pos, para_element in reversed(preceding_text_paras[-4:]):  # Look at last 4 paras
            # Extract text from this paragraph
            text_elems = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            text = ''.join([t.text for t in text_elems if t.text]).strip()

            if not text:
                continue  # Skip empty paragraphs

            # Don't use text that looks like a section heading
            if self._classify_heading_from_text(text, para_element):
                break  # Stop searching, we hit a heading

            # Don't use text that's too long (likely regular content)
            # Captions are typically short and descriptive
            if len(text) > 250:
                continue

            # Don't use bullet points or numbered lists
            if text.startswith('←') or text.startswith('-'):
                continue

            # This looks like a caption!
            candidates.append(text)

        # Return the most relevant caption (usually the last non-empty, non-heading text)
        if candidates:
            return candidates[-1]  # Most recent caption before the table

        return None

    def _detect_chart_positions(self) -> List[Dict[str, Any]]:
        """
        Detect chart positions in document by finding all visual content.

        Uses multiple detection methods to find charts/images/shapes:
        - DrawingML elements (native Office charts/images)
        - Alternative content (legacy Office format)
        - Embedded objects and shapes
        - Tables that contain charts

        Positions are tracked in the same scale as sections (element count).

        Handles multiple charts in the same paragraph by counting all inline
        and anchored shapes.

        Automatically skips charts that appear before main content (cover/TOC pages).

        Returns:
            List of chart positions with proper positioning
        """
        chart_positions = []
        element_position = 0
        chart_count = 0

        # Traverse document body in actual order, incrementing position for EVERY element
        for element in self.doc.element.body:
            visual_count = 0

            if element.tag.endswith('p'):
                # Count ALL visual elements in paragraph (inline + anchored)
                visual_count = self._count_visual_elements_in_paragraph(element)

            elif element.tag.endswith('tbl'):
                # Check if table contains visual content
                if self._element_has_visual_content(element):
                    visual_count = 1

            # Record a position for EACH visual element found
            for i in range(visual_count):
                # Skip charts before main content starts (cover page, TOC)
                if hasattr(self, '_main_content_start_pos') and element_position < self._main_content_start_pos:
                    print(f"  📊 Chart {chart_count} at position {element_position} SKIPPED (before main content)")
                    chart_count += 1
                    continue

                chart_positions.append({
                    'position': element_position,  # Position of this element
                    'index': chart_count,
                    'assigned_section': None
                })
                print(f"  📊 Chart {chart_count} detected at position {element_position} (shape {i+1}/{visual_count})")
                chart_count += 1

            # ALWAYS increment position to stay in sync with section position tracking
            element_position += 1

        return chart_positions

    def _count_visual_elements_in_paragraph(self, para_elem) -> int:
        """Count all visual elements (inline shapes + anchored shapes) in a paragraph."""
        ns_main = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        ns_wp = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'

        count = 0

        # Count inline shapes (wp:inline)
        inline_shapes = para_elem.findall(f'.//{ns_wp}inline')
        count += len(inline_shapes)
        print(f"    → Found {len(inline_shapes)} inline shapes")

        # Count anchored shapes (wp:anchor)
        anchored_shapes = para_elem.findall(f'.//{ns_wp}anchor')
        count += len(anchored_shapes)
        print(f"    → Found {len(anchored_shapes)} anchored shapes")

        return count

    def _paragraph_has_visual_content(self, para_elem) -> bool:
        """Check if paragraph contains any visual content (drawings, images, shapes)."""
        # Namespace constants
        ns_main = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        ns_wp = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
        ns_a = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
        ns_pic = '{http://schemas.openxmlformats.org/drawingml/2006/picture}'

        # Check for drawing elements in runs (most common for images/charts)
        for run in para_elem.findall(f'.//{ns_main}r'):
            drawing = run.find(f'.//{ns_main}drawing')
            if drawing is not None:
                return True

            # Check for pict (legacy VML format)
            pict = run.find(f'.//{ns_main}pict')
            if pict is not None:
                return True

        # Check for alternate content (used for compatibility between formats)
        if para_elem.find(f'.//{ns_main}AlternateContent') is not None:
            return True

        # Check for embedded objects
        if para_elem.find(f'.//{ns_main}object') is not None:
            return True

        # Check for bookmarks or special content
        if para_elem.find(f'.//{ns_main}bookmarkStart') is not None:
            # Might indicate a special element like a chart
            pass

        return False

    def _element_has_visual_content(self, elem) -> bool:
        """Check if any element contains visual content (charts, images, shapes, etc.)."""
        ns_main = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        ns_wp = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'

        # Check for modern drawingML elements
        if elem.find(f'.//{ns_main}drawing') is not None:
            return True

        # Check for alternate content (compatibility format)
        if elem.find(f'.//{ns_main}AlternateContent') is not None:
            return True

        # Check for embedded objects
        if elem.find(f'.//{ns_main}object') is not None:
            return True

        # Check for legacy VML format (older Office documents)
        if elem.find(f'.//{ns_main}pict') is not None:
            return True

        # Check for shapes
        if elem.find(f'.//{ns_main}wps') is not None:  # Word Processing Shape
            return True

        return False

    def _assign_elements_to_hierarchy(
        self,
        main_sections: List[Dict],
        tables: List[Dict],
        charts: List[Dict]
    ) -> None:
        """
        Assign tables and charts to sections/subsections using proximity-based assignment.

        Each element (table/chart) is assigned to the closest preceding section (main or subsection).
        No complex scoring — just find what comes immediately before.
        """
        if not main_sections:
            return

        # Build sorted list of all section positions
        section_positions = []
        position_to_section = {}

        for main_sec in main_sections:
            pos = main_sec['position']
            section_positions.append(pos)
            position_to_section[pos] = (main_sec, True)  # (section, is_main)

            for subsec in main_sec['subsections']:
                pos = subsec['position']
                section_positions.append(pos)
                position_to_section[pos] = (subsec, False)

        section_positions.sort()

        def find_closest_preceding_section(element_pos):
            """Find the section closest before this element position."""
            closest_pos = None

            for sec_pos in section_positions:
                if sec_pos < element_pos:
                    closest_pos = sec_pos
                else:
                    break  # Stop at first section at or after element

            if closest_pos is not None:
                section, is_main = position_to_section[closest_pos]
                return section, is_main

            return None, None

        # Assign tables
        for table_info in tables:
            table_pos = table_info['position']
            section, is_main = find_closest_preceding_section(table_pos)

            if section:
                section['tables'].append(table_info)
                table_info['assigned_section'] = section['id']
                level_type = "Sub" if not is_main else "Main"
                print(f"  📋 Table {table_info['index']} at pos {table_pos} → {section['title_ar']} ({level_type})")
            elif main_sections:
                # Fallback: assign to first section if nothing precedes
                main_sections[0]['tables'].append(table_info)
                table_info['assigned_section'] = main_sections[0]['id']
                print(f"  📋 Table {table_info['index']} at pos {table_pos} → {main_sections[0]['title_ar']} (FALLBACK)")

        # Assign charts using same approach
        for chart_info in charts:
            chart_pos = chart_info['position']
            section, is_main = find_closest_preceding_section(chart_pos)

            if section:
                section['charts'].append(chart_info)
                chart_info['assigned_section'] = section['id']
                level_type = "Sub" if not is_main else "Main"
                print(f"  📊 Chart {chart_info['index']} at pos {chart_pos} → {section['title_ar']} ({level_type})")
            elif main_sections:
                # Fallback: assign to first section if nothing precedes
                main_sections[0]['charts'].append(chart_info)
                chart_info['assigned_section'] = main_sections[0]['id']
                print(f"  📊 Chart {chart_info['index']} at pos {chart_pos} → {main_sections[0]['title_ar']} (FALLBACK)")

    def _classify_heading(self, para) -> str:
        """
        Classify a paragraph as 'main' section, 'sub' section, or None.

        Uses Word heading styles as primary method, with pattern fallbacks.
        This is more reliable than font size which can vary.

        Returns:
            'main' - Main section (Heading 1, or أولاً, ثانياً, etc.)
            'sub' - Subsection (Heading 2, or numbered like 2.1, 2.2)
            None - Not a heading
        """
        text = para.text.strip()
        if not text:
            return None

        # Skip if text is too long (likely paragraph content)
        if len(text) > 300:
            return None

        # Skip table of contents entries (small font in TOC vs larger in actual content)
        # TOC entries are typically 9-10pt, actual headings are 13-16pt
        if para.runs and para.runs[0].font.size:
            font_size = para.runs[0].font.size.pt
            if font_size < 11:  # Skip TOC entries
                return None

        # PRIMARY METHOD: Check Word heading styles (most reliable)
        if para.style and 'Heading' in para.style.name:
            # Remove spaces for comparison (handles both "Heading 1" and "Heading1")
            normalized = para.style.name.replace(' ', '')
            if 'Heading1' in normalized:
                return 'main'
            elif 'Heading' in normalized:
                return 'sub'

        # FALLBACK: Check for main section patterns (Arabic ordinals with colon)
        # This catches cases where heading styles aren't applied consistently
        main_patterns = [
            r'^(أولاً|ثانياً|ثالثاً|رابعاً|خامساً|سادساً|سابعاً|ثامناً|تاسعاً|عاشراً):',
        ]

        for pattern in main_patterns:
            if re.search(pattern, text):
                return 'main'

        # FALLBACK: Check for subsection patterns (numbered like 2.1, 2.2)
        # Must start with digit.digit followed by space
        sub_patterns = [
            r'^\d+\.\d+\s',  # Matches: "2.1 ", "3.2 ", etc.
        ]

        for pattern in sub_patterns:
            if re.search(pattern, text):
                return 'sub'

        return None

    def _get_heading_level(self, para) -> int:
        """Get heading level (1-9) from paragraph style."""
        if para.style and 'Heading' in para.style.name:
            match = re.search(r'Heading (\d+)', para.style.name)
            if match:
                return int(match.group(1))

        # Estimate level based on font size
        if para.runs and para.runs[0].font.size:
            size_pt = para.runs[0].font.size.pt
            if size_pt >= 18:
                return 1
            elif size_pt >= 14:
                return 2
            else:
                return 3

        return 2  # Default

    def _generate_section_id(self, title: str, index: int) -> str:
        """
        Generate a unique ID for a section.

        Uses transliteration of Arabic or simple slugification.
        """
        # Simple approach: use index and first few words
        words = title.strip().split()[:3]
        slug = '_'.join(words).lower()

        # Remove special characters
        slug = re.sub(r'[^\w\s_-]', '', slug)
        slug = re.sub(r'[-\s]+', '_', slug)

        return f"section_{index}_{slug}"[:50]  # Limit length

    def _translate_section_title(self, title_ar: str) -> str:
        """
        Simple translation of common Arabic section titles to English.

        For unknown titles, returns transliteration or original.
        """
        translations = {
            'الملخص التنفيذي': 'Executive Summary',
            'المقدمة': 'Introduction',
            'الخلاصة': 'Conclusion',
            'التوصيات': 'Recommendations',
            'حالات الاستخدام': 'Use Cases',
            'خارطة الطريق': 'Roadmap',
        }

        # Check for exact matches
        title_clean = title_ar.strip()
        if title_clean in translations:
            return translations[title_clean]

        # Check for partial matches
        for ar_key, en_val in translations.items():
            if ar_key in title_clean:
                return en_val

        # Check for "Analysis" pattern
        if 'تحليل' in title_clean or 'التحليل' in title_clean:
            # Extract number if present
            match = re.search(r'(الأول|الثاني|الثالث|الرابع|الخامس|\d+)', title_clean)
            if match:
                num_map = {
                    'الأول': 'First',
                    'الثاني': 'Second',
                    'الثالث': 'Third',
                    'الرابع': 'Fourth',
                    'الخامس': 'Fifth'
                }
                num_text = num_map.get(match.group(1), match.group(1))
                return f"Analysis {num_text}"
            return "Analysis"

        # Return original if no translation found
        return title_ar

    def _find_table_position(self, table, table_idx: int) -> int:
        """
        Find the actual position of a table in the document.

        Uses a more sophisticated approach that:
        1. Finds the closest preceding paragraph/heading
        2. Counts elements from start of document
        3. Returns accurate position for proximity matching

        Returns: The cumulative position of the table in the document
        """
        # Count ALL elements in document body (matches _detect_all_headings counting)
        # This is critical for consistent position tracking across all elements
        element_count = 0
        table_count = 0

        for element in self.doc.element.body:
            # Check if this is a table element
            if element.tag.endswith('tbl'):
                if table_count == table_idx:
                    # Found our table - return current element count
                    # This gives us the position in the overall document flow
                    return element_count
                table_count += 1

            # ALWAYS increment position for EVERY element, not just p and tbl
            # This matches heading position counting and ensures consistency
            element_count += 1

        # If table not found, return approximate position
        # This should rarely happen
        total_paras = len(self.doc.paragraphs)
        total_tables = len(self.doc.tables)
        return total_paras + total_tables // 2

    def _extract_table_data(self, table) -> Dict[str, Any]:
        """
        Extract complete table data with columns and rows.

        Returns:
            Dictionary with columns, rows, and metadata
        """
        if len(table.rows) == 0:
            return {'columns': [], 'rows': []}

        # Get column headers from first row
        columns = [cell.text.strip() for cell in table.rows[0].cells]

        # Get data rows
        rows = []
        for row in table.rows[1:]:
            row_data = {}
            cells = [cell.text.strip() for cell in row.cells]

            for col_idx, col_name in enumerate(columns):
                if col_idx < len(cells):
                    row_data[col_name] = cells[col_idx]
                else:
                    row_data[col_name] = ''

            rows.append(row_data)

        return {
            'columns': columns,
            'rows': rows,
            'row_count': len(rows),
            'col_count': len(columns)
        }

    def _extract_metadata(self) -> Dict[str, Any]:
        """Extract document metadata."""
        core_props = self.doc.core_properties

        return {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else '',
            'total_paragraphs': len(self.doc.paragraphs),
            'total_tables': len(self.doc.tables)
        }
