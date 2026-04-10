"""
WordBuilder — the main public interface for sword-word-builder.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any, Literal, Self

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from .config import (
    DocumentConfig, TextStyle, TableStyle, ChartStyle,
    PAGE_SIZES, hex_to_rgb,
)
from ._xml_utils import (
    apply_run_formatting,
    set_paragraph_rtl,
    set_paragraph_spacing,
    add_paragraph_bottom_border,
    add_paragraph_top_border,
    add_page_number_field,
    add_page_x_of_y_field,
    add_text_and_page_number_field,
    add_hyperlink as _add_hyperlink_to_para,
    alignment_const,
    is_arabic,
)
from ._table_builder import build_table
from ._chart_injector import (
    _PendingChart, inject_charts, sentinel_name,
)
from ._markdown import MarkdownRenderer
from ._cover_page import CoverPage


class WordBuilder:
    """
    Fluent Word document builder.

    Usage::

        from sword_word_builder import WordBuilder, DocumentConfig

        builder = WordBuilder(DocumentConfig(default_font="Calibri"))
        builder.add_heading("Report Title", level=1)
        builder.add_paragraph("Introduction text.")
        builder.add_table(data, headers=["Name", "Score"])
        builder.add_chart(chart_data, chart_type="column")
        builder.save("report.docx")

    All ``add_*`` methods return ``self`` to allow optional method chaining.
    """

    def __init__(self, config: DocumentConfig | None = None) -> None:
        self._config = config or DocumentConfig()
        self._doc = Document()
        self._pending_charts: list[_PendingChart] = []
        self._doc_pr_counter = 1   # unique IDs for drawings (images + charts)
        self._markdown_renderer = MarkdownRenderer(self._config)
        self._apply_document_config()

    # ------------------------------------------------------------------
    # Internal: document configuration
    # ------------------------------------------------------------------

    def _apply_document_config(self) -> None:
        cfg = self._config
        doc = self._doc

        # Page geometry
        section = doc.sections[0]
        page_w_emu, page_h_emu = PAGE_SIZES[cfg.page_size]
        section.page_width  = int(page_w_emu)
        section.page_height = int(page_h_emu)
        section.top_margin    = Cm(cfg.margin_top)
        section.bottom_margin = Cm(cfg.margin_bottom)
        section.left_margin   = Cm(cfg.margin_left)
        section.right_margin  = Cm(cfg.margin_right)

        # Normal paragraph style
        normal = doc.styles["Normal"]
        normal.font.name = cfg.default_font
        normal.font.size = Pt(cfg.default_font_size)
        _set_cs_font(normal, cfg.default_font)
        normal.paragraph_format.space_after = Pt(cfg.default_paragraph_space_after)
        if cfg.line_spacing is not None:
            normal.paragraph_format.line_spacing = Pt(cfg.line_spacing)

        # Heading styles
        heading_font = cfg.heading_font or cfg.default_font
        _sizes = {1: cfg.heading1_size, 2: cfg.heading2_size, 3: cfg.heading3_size}
        _colors = {
            1: cfg.accent_color,
            2: cfg.secondary_color,
            3: cfg.accent_color,
        }
        for level in (1, 2, 3):
            try:
                h = doc.styles[f"Heading {level}"]
                h.font.name = heading_font
                h.font.size = Pt(_sizes[level])
                h.font.bold = cfg.heading_bold
                r, g, b = hex_to_rgb(_colors[level])
                h.font.color.rgb = RGBColor(r, g, b)
                _set_cs_font(h, heading_font)
            except Exception:
                pass

        # Apply default header/footer from config
        self._apply_header_footer_from_config(section)

    def _apply_header_footer_from_config(self, section) -> None:
        cfg = self._config
        if cfg.skip_first_page_header_footer:
            section.different_first_page_header_footer = True
        self._write_header(section, cfg.header_type, cfg.header_text,
                           cfg.header_image_path, cfg.header_image_height_cm,
                           cfg.header_alignment)
        self._write_footer(section, cfg.footer_type, cfg.footer_text,
                           cfg.footer_image_path, cfg.footer_image_height_cm,
                           cfg.footer_alignment)

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def add_cover_page(self, cover: CoverPage) -> Self:
        """Add a cover page followed by a page break.

        ``cover`` is a :class:`CoverPage` canvas object that holds any
        combination of content components (headings, paragraphs, images,
        tables, charts, separators, etc.).  All components are rendered
        onto this builder instance in the order they were added, then a
        page break is appended automatically.

        Example::

            cover = CoverPage()
            cover.add_spacer(5)
            cover.add_heading("Annual Report 2025", level=1)
            cover.add_horizontal_separator(native=True)
            cover.add_paragraph(
                "Strategy Division",
                style=TextStyle(size=14, italic=True, alignment="CENTER"),
            )
            builder.add_cover_page(cover)

        For the classic fixed layout use :meth:`CoverPage.preset`::

            builder.add_cover_page(
                CoverPage.preset("Title", subtitle="Sub", metadata={"Author": "Me"})
            )
        """
        cover._render(self)

        # Page break
        last_para = self._doc.add_paragraph()
        br_run = last_para.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        br_run._r.append(br)

        return self

    # ------------------------------------------------------------------
    # Text content
    # ------------------------------------------------------------------

    def add_heading(
        self,
        text: str,
        level: int = 1,
        style: TextStyle | None = None,
        rtl: bool | None = None,
        separator: bool | None = None,
    ) -> Self:
        """Add a heading at the given level (1–6)."""
        cfg = self._config
        rtl = _resolve_rtl(rtl, cfg)

        try:
            para = self._doc.add_heading(level=level)
        except Exception:
            para = self._doc.add_paragraph(style="Heading 1")
        para.clear()

        if rtl:
            set_paragraph_rtl(para, True)
        # Default heading alignment: paragraph start (physically right for RTL)
        para.alignment = _resolve_paragraph_alignment(None, rtl)

        run = para.add_run(text)
        font_name = cfg.heading_font or cfg.default_font
        size_map = {1: cfg.heading1_size, 2: cfg.heading2_size, 3: cfg.heading3_size}
        font_size = size_map.get(level, cfg.default_font_size + 2)
        color = cfg.accent_color if level != 2 else cfg.secondary_color

        apply_run_formatting(
            run,
            font_name=font_name,
            font_size_pt=font_size,
            bold=cfg.heading_bold,
            color_hex=color,
            rtl=rtl,
        )

        # Apply TextStyle overrides
        if style:
            _apply_text_style_to_run(run, style)
            if style.space_before is not None or style.space_after is not None:
                set_paragraph_spacing(para, style.space_before, style.space_after)
            if style.alignment != "LEFT":
                para.alignment = _resolve_paragraph_alignment(style.alignment, rtl)

        # Separator
        sep_defaults = {1: cfg.heading1_separator, 2: cfg.heading2_separator,
                        3: cfg.heading3_separator}
        should_sep = separator if separator is not None else sep_defaults.get(level, False)
        if should_sep:
            add_paragraph_bottom_border(para)
            set_paragraph_spacing(para, space_after_pt=16)

        return self

    def add_paragraph(
        self,
        text: str,
        style: TextStyle | None = None,
        rtl: bool | None = None,
        bold: bool = False,
        italic: bool = False,
        alignment: str | None = None,
        space_after: float | None = None,
        keep_with_next: bool = False,
    ) -> Self:
        """Add a styled paragraph."""
        cfg = self._config
        rtl = _resolve_rtl(rtl, cfg)

        para = self._doc.add_paragraph()
        if rtl:
            set_paragraph_rtl(para, True)
        al = alignment or (style.alignment if style else None)
        para.alignment = _resolve_paragraph_alignment(al, rtl)

        sa = space_after if space_after is not None else (
            style.space_after if style and style.space_after is not None
            else cfg.default_paragraph_space_after
        )
        set_paragraph_spacing(para, space_after_pt=sa)

        if keep_with_next or (style and style.keep_with_next):
            para.paragraph_format.keep_with_next = True

        run = para.add_run(text)
        apply_run_formatting(
            run,
            font_name=cfg.default_font,
            font_size_pt=cfg.default_font_size,
            bold=bold or (style.bold if style else False),
            italic=italic or (style.italic if style else False),
            underline=(style.underline if style else False),
            color_hex=(style.color if style else None) or cfg.body_color,
            rtl=rtl,
        )
        if style:
            _apply_text_style_to_run(run, style)

        return self

    def add_hyperlink(
        self,
        url: str,
        display_text: str,
        paragraph=None,
        rtl: bool | None = None,
    ) -> Self:
        """Add a hyperlink. If paragraph is None, a new paragraph is created."""
        cfg = self._config

        # Auto-detect RTL from Arabic text BEFORE resolving defaults,
        # so Arabic display text always gets RTL unless explicitly overridden.
        if rtl is None:
            link_rtl = is_arabic(display_text)
        else:
            link_rtl = rtl

        para_rtl = _resolve_rtl(rtl, cfg)
        if paragraph is None:
            paragraph = self._doc.add_paragraph()
            if para_rtl:
                set_paragraph_rtl(paragraph, True)

        _add_hyperlink_to_para(
            paragraph, url, display_text,
            font_name=cfg.default_font,
            font_size_pt=cfg.default_font_size,
            rtl=link_rtl,
        )
        return self


    def add_markdown(
        self,
        markdown_text: str,
        rtl: bool | None = None,
        base_style: TextStyle | None = None,
    ) -> Self:
        """Render markdown text into the document."""
        rtl = _resolve_rtl(rtl, self._config)
        renderer = MarkdownRenderer(self._config, base_style)
        renderer.render(self._doc, markdown_text, rtl)
        return self

    def add_horizontal_separator(
        self,
        color: str | None = None,
        native: bool = False,
    ) -> Self:
        """Add a visual horizontal separator line.

        Args:
            color: Hex color string (with or without '#').  Defaults to
                   ``DocumentConfig.body_color``.
            native: When ``True``, emits a real Word paragraph border
                    (``<w:pBdr><w:bottom …/>``) instead of 50 Unicode ``─``
                    characters.  Native borders scale correctly at any zoom
                    level and look cleaner in print and PDF export.

        Examples::

            builder.add_horizontal_separator()                       # Unicode (default)
            builder.add_horizontal_separator(native=True)            # Word border
            builder.add_horizontal_separator(native=True, color="2E74B5")
        """
        cfg = self._config
        resolved_color = color or cfg.body_color
        para = self._doc.add_paragraph()
        set_paragraph_spacing(para, space_before_pt=8, space_after_pt=8)
        if native:
            add_paragraph_bottom_border(para, color_hex=resolved_color)
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("─" * 50)
            apply_run_formatting(
                run,
                font_name=cfg.default_font,
                font_size_pt=8,
                color_hex=resolved_color,
            )
        return self

    def add_banner(
        self,
        text: str,
        bg_color: str = "2E74B5",
        text_color: str = "FFFFFF",
        font_size: int = 14,
        bold: bool = True,
        rtl: bool | None = None,
    ) -> Self:
        """Add a full-width single-row coloured banner, typically used as a
        section header.

        This is a convenience wrapper around :meth:`add_table` that produces a
        borderless, single-cell, single-row table spanning the full page width::

            builder.add_banner("أولاً  التحليل الديموغرافي", bg_color="0D2D5E", text_color="B89A00")

        Args:
            text: The banner label text.
            bg_color: Background hex colour (default: accent blue ``2E74B5``).
            text_color: Text hex colour (default: white ``FFFFFF``).
            font_size: Font size in pt (default: 14).
            bold: Whether the text is bold (default: True).
            rtl: Override RTL direction; ``None`` inherits from ``DocumentConfig``.
        """
        style = TableStyle(
            header_bg_color=bg_color,
            header_text_color=text_color,
            header_bold=bold,
            header_font_size=font_size,
            show_inner_borders=False,
            show_outer_border=False,
            alt_row_bg_color=None,
        )
        return self.add_table([], headers=[text], style=style, rtl=rtl)

    def add_page_break(self) -> Self:
        """Insert a page break."""
        para = self._doc.add_paragraph()
        run = para.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)
        return self

    def add_spacer(self, height_pt: float = 12.0) -> Self:
        """Insert an empty paragraph with the given space_after."""
        para = self._doc.add_paragraph()
        set_paragraph_spacing(para, space_before_pt=0, space_after_pt=height_pt)
        return self

    def add_picture(
        self,
        image_path: str | Path,
        width_cm: float | None = None,
        height_cm: float | None = None,
        alignment: str = "CENTER",
        caption: str | None = None,
    ) -> Self:
        """Insert an image."""
        para = self._doc.add_paragraph()
        para.alignment = alignment_const(alignment)

        run = para.add_run()
        kwargs: dict[str, Any] = {}
        if width_cm is not None:
            kwargs["width"] = Cm(width_cm)
        if height_cm is not None:
            kwargs["height"] = Cm(height_cm)
        run.add_picture(str(image_path), **kwargs)
        self._doc_pr_counter += 1

        if caption:
            cap_para = self._doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(caption)
            apply_run_formatting(
                cap_run,
                font_name=self._config.default_font,
                font_size_pt=self._config.default_font_size - 1,
                italic=True,
                color_hex=self._config.body_color,
                rtl=is_arabic(caption),
            )
        return self

    # ------------------------------------------------------------------
    # Tables
    # ------------------------------------------------------------------

    def add_table(
        self,
        data: list[dict] | list[list] | list[tuple],
        headers: list[str] | None = None,
        style: TableStyle | None = None,
        rtl: bool | None = None,
        caption: str | None = None,
        table_type: Literal["data", "metrics"] = "data",
    ) -> Self:
        """Add a table.

        data formats:
          - list[dict]  → headers auto-extracted from keys
          - list[list]  → requires headers=
          - list[tuple] → for table_type="metrics" (label-value pairs)
          - dict        → for table_type="metrics"

        Cell values may be wrapped as ``(value, CellStyle)`` tuples to apply
        per-cell background, text colour, bold, or font-size overrides::

            from sword_word_builder import CellStyle

            builder.add_table(
                [{"Status": ("Pass ✔", CellStyle(bg_color="1E7B4E", text_color="FFFFFF", bold=True))}]
            )

        For ``table_type="metrics"``, supply ``headers=["LabelCol", "ValueCol"]``
        to override the default ``["Label", "Value"]`` column names.
        """
        cfg = self._config
        rtl = _resolve_rtl(rtl, cfg)
        tbl_style = style or TableStyle()

        build_table(self._doc, data, headers, tbl_style, cfg, rtl, table_type)

        # Spacing after table
        self._doc.add_paragraph()

        if caption:
            cap_para = self._doc.add_paragraph()
            if rtl:
                set_paragraph_rtl(cap_para, True)
            else:
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(caption)
            apply_run_formatting(
                cap_run,
                font_name=cfg.default_font,
                font_size_pt=cfg.default_font_size - 1,
                italic=True,
                color_hex=cfg.body_color,
                rtl=rtl or is_arabic(caption),
            )
        return self

    # ------------------------------------------------------------------
    # Charts
    # ------------------------------------------------------------------

    def add_chart(
        self,
        data: dict,
        chart_type: Literal["column", "bar", "line", "pie"] = "column",
        style: ChartStyle | None = None,
        rtl: bool | None = None,
    ) -> Self:
        """Add a native Word chart (editable in Word with embedded Excel data).

        data format::

            {
                "title": "My Chart",           # optional
                "categories": ["Q1", "Q2", "Q3"],
                "series": [
                    {"name": "Revenue", "values": [100, 150, 130], "color": "4472C4"},
                ],
            }

        chart_type: "column" (vertical bars), "bar" (horizontal bars), "line", "pie"
        """
        chart_style = style or ChartStyle()
        idx = len(self._pending_charts)

        # Insert sentinel paragraph with bookmark
        para = self._doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Embed a w:bookmarkStart so we can find this paragraph after saving
        bk_start = OxmlElement("w:bookmarkStart")
        bk_start.set(qn("w:id"), str(idx))
        bk_start.set(qn("w:name"), sentinel_name(idx))
        bk_end = OxmlElement("w:bookmarkEnd")
        bk_end.set(qn("w:id"), str(idx))
        para._p.insert(0, bk_start)
        para._p.append(bk_end)

        self._pending_charts.append(_PendingChart(
            index=idx,
            data=data,
            chart_type=chart_type,
            style=chart_style,
        ))

        # Spacing after chart
        self._doc.add_paragraph()
        return self

    # ------------------------------------------------------------------
    # Document structure
    # ------------------------------------------------------------------

    def add_section(
        self,
        orientation: Literal["portrait", "landscape"] = "portrait",
    ) -> Self:
        """Add a new document section (section break)."""
        new_section = self._doc.add_section()
        cfg = self._config
        page_w_emu, page_h_emu = PAGE_SIZES[cfg.page_size]

        if orientation == "landscape":
            new_section.orientation = WD_ORIENT.LANDSCAPE
            new_section.page_width  = int(page_h_emu)
            new_section.page_height = int(page_w_emu)
        else:
            new_section.orientation = WD_ORIENT.PORTRAIT
            new_section.page_width  = int(page_w_emu)
            new_section.page_height = int(page_h_emu)

        new_section.top_margin    = Cm(cfg.margin_top)
        new_section.bottom_margin = Cm(cfg.margin_bottom)
        new_section.left_margin   = Cm(cfg.margin_left)
        new_section.right_margin  = Cm(cfg.margin_right)
        return self

    def set_header(
        self,
        header_type: Literal["none", "page_number", "text", "image", "text_and_page_number"] = "none",
        text: str = "",
        image_path: str | Path | None = None,
        image_height_cm: float = 1.5,
        alignment: Literal["LEFT", "CENTER", "RIGHT"] = "CENTER",
        skip_first_page: bool | None = None,
    ) -> Self:
        """Configure the document header (overrides DocumentConfig defaults)."""
        section = self._doc.sections[0]
        if skip_first_page is not None:
            section.different_first_page_header_footer = skip_first_page
        self._write_header(section, header_type, text,
                           str(image_path) if image_path else None,
                           image_height_cm, alignment)
        return self

    def set_footer(
        self,
        footer_type: Literal["none", "page_number", "text", "image", "text_and_page_number"] = "page_number",
        text: str = "",
        image_path: str | Path | None = None,
        image_height_cm: float = 1.0,
        alignment: Literal["LEFT", "CENTER", "RIGHT"] = "CENTER",
        skip_first_page: bool | None = None,
    ) -> Self:
        """Configure the document footer (overrides DocumentConfig defaults)."""
        section = self._doc.sections[0]
        if skip_first_page is not None:
            section.different_first_page_header_footer = skip_first_page
        self._write_footer(section, footer_type, text,
                           str(image_path) if image_path else None,
                           image_height_cm, alignment)
        return self

    # ------------------------------------------------------------------
    # Header/footer writers
    # ------------------------------------------------------------------

    def _write_header(self, section, htype: str, text: str,
                       image_path: str | None, image_height_cm: float,
                       alignment: str) -> None:
        cfg = self._config
        header = section.header
        # Clear existing content
        for para in header.paragraphs:
            para.clear()

        if htype == "none":
            return

        rtl = is_arabic(text) if text else cfg.default_rtl
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        if rtl:
            set_paragraph_rtl(para, True)
        para.alignment = _resolve_paragraph_alignment(alignment, rtl)

        font_size = cfg.header_font_size or cfg.default_font_size
        text_color = cfg.header_text_color or cfg.body_color

        if htype == "page_number":
            add_page_x_of_y_field(para)
        elif htype == "text":
            run = para.add_run(text)
            apply_run_formatting(run, font_name=cfg.default_font,
                                 font_size_pt=font_size,
                                 color_hex=text_color,
                                 rtl=rtl)
        elif htype == "text_and_page_number":
            add_text_and_page_number_field(
                para, prefix=text,
                font_name=cfg.default_font,
                font_size_pt=font_size,
                rtl=rtl,
            )
        elif htype == "image" and image_path:
            _insert_stretched_image(para, image_path, image_height_cm, section)

        if cfg.header_bottom_border_color:
            add_paragraph_bottom_border(para, cfg.header_bottom_border_color)
        if cfg.header_spacing_after:
            set_paragraph_spacing(para, space_after_pt=cfg.header_spacing_after)

    def _write_footer(self, section, ftype: str, text: str,
                       image_path: str | None, image_height_cm: float,
                       alignment: str) -> None:
        cfg = self._config
        footer = section.footer
        for para in footer.paragraphs:
            para.clear()

        if ftype == "none":
            return

        rtl = is_arabic(text) if text else cfg.default_rtl
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        if rtl:
            set_paragraph_rtl(para, True)
        para.alignment = _resolve_paragraph_alignment(alignment, rtl)

        font_size = cfg.footer_font_size or cfg.default_font_size
        text_color = cfg.footer_text_color or cfg.body_color

        if ftype == "page_number":
            add_page_x_of_y_field(para)
        elif ftype == "text":
            run = para.add_run(text)
            apply_run_formatting(run, font_name=cfg.default_font,
                                 font_size_pt=font_size,
                                 color_hex=text_color,
                                 rtl=rtl)
        elif ftype == "text_and_page_number":
            add_text_and_page_number_field(
                para, prefix=text,
                font_name=cfg.default_font,
                font_size_pt=font_size,
                rtl=rtl,
            )
        elif ftype == "image" and image_path:
            _insert_stretched_image(para, image_path, image_height_cm, section)

        if cfg.footer_top_border_color:
            add_paragraph_top_border(para, cfg.footer_top_border_color)
        if cfg.footer_spacing_before:
            set_paragraph_spacing(para, space_before_pt=cfg.footer_spacing_before)

    # ------------------------------------------------------------------
    # Output
    # ------------------------------------------------------------------

    def build(self) -> io.BytesIO:
        """Build and return the document as a BytesIO buffer."""
        buf = io.BytesIO()
        self._doc.save(buf)
        buf.seek(0)

        if self._pending_charts:
            buf = inject_charts(buf, self._pending_charts,
                                 doc_pr_id_start=self._doc_pr_counter + 100)
        return buf

    def save(self, path: str | Path) -> None:
        """Build and save the document to a file path."""
        buf = self.build()
        with open(path, "wb") as f:
            f.write(buf.read())


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _resolve_rtl(rtl: bool | None, config: DocumentConfig) -> bool:
    return config.default_rtl if rtl is None else rtl


def _resolve_paragraph_alignment(
    alignment: str | None,
    rtl: bool,
) -> WD_ALIGN_PARAGRAPH:
    """Map a logical alignment string to the correct OOXML ``jc`` value.

    OOXML §17.18.44: when ``w:bidi=1`` is set on a paragraph, ``jc=left``
    renders as *physically right* (the RTL reading-start side) and ``jc=right``
    renders as *physically left* (the RTL reading-end side).  CENTER and
    JUSTIFY are direction-neutral and require no remapping.

    Mapping when ``rtl=True``:

    +-----------+------------------+---------------------+
    | Requested | jc value emitted | Physical appearance |
    +===========+==================+=====================+
    | None / RIGHT | left          | physically right ✓  |
    | LEFT         | right         | physically left ✓   |
    | CENTER       | center        | centred ✓           |
    | JUSTIFY      | both          | justified ✓         |
    +-----------+------------------+---------------------+
    """
    if alignment is None:
        # Default: paragraph start — jc=left works for both LTR and RTL
        return WD_ALIGN_PARAGRAPH.LEFT

    al = alignment.upper()

    if rtl:
        if al == "CENTER":
            return WD_ALIGN_PARAGRAPH.CENTER
        if al in ("JUSTIFY", "BOTH"):
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        if al == "LEFT":
            # Physically left = reading end in RTL → needs jc=right
            return WD_ALIGN_PARAGRAPH.RIGHT
        # RIGHT (or unrecognised) → physically right = reading start → jc=left
        return WD_ALIGN_PARAGRAPH.LEFT

    return alignment_const(alignment)


def _apply_text_style_to_run(run, style: TextStyle) -> None:
    """Apply TextStyle overrides on top of existing run formatting."""
    if style.font is not None:
        run.font.name = style.font
    if style.size is not None:
        run.font.size = Pt(style.size)
    if style.bold:
        run.font.bold = True
    if style.italic:
        run.font.italic = True
    if style.underline:
        run.font.underline = True
    if style.color is not None:
        r, g, b = hex_to_rgb(style.color)
        run.font.color.rgb = RGBColor(r, g, b)


def _set_cs_font(style_obj, font_name: str) -> None:
    """Set complex-script font on a Word style object."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        element = style_obj.element
        rPr = element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            element.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:cs"), font_name)
    except Exception:
        pass


def _insert_stretched_image(para, image_path: str, height_cm: float, section) -> None:
    """Insert an image into a header/footer paragraph, stretching it to page width."""
    run = para.add_run()
    try:
        run.add_picture(image_path, width=section.page_width)
        # Extend paragraph to reach page edge (negative indents)
        fmt = para.paragraph_format
        fmt.left_indent  = -section.left_margin
        fmt.right_indent = -section.right_margin
    except Exception:
        pass
