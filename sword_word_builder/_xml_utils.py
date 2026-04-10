"""
Internal XML helpers for sword-word-builder.

All functions operate on python-docx objects and mutate their underlying lxml XML.
Nothing in this module is part of the public API.
"""
from __future__ import annotations

import re
from typing import TYPE_CHECKING

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.table import _Cell

# ---------------------------------------------------------------------------
# Alignment mapping
# ---------------------------------------------------------------------------

_ALIGN_MAP = {
    "LEFT":    WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER":  WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT":   WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def alignment_const(alignment: str) -> WD_ALIGN_PARAGRAPH:
    return _ALIGN_MAP.get(alignment.upper(), WD_ALIGN_PARAGRAPH.LEFT)


# ---------------------------------------------------------------------------
# Paragraph-level helpers
# ---------------------------------------------------------------------------

def set_paragraph_rtl(paragraph: "Paragraph", rtl: bool) -> None:
    """Insert or remove <w:bidi/> in pPr.

    This function only manages the directionality flag; it does NOT touch
    paragraph alignment (jc).  Callers are responsible for setting alignment
    via :func:`_resolve_paragraph_alignment` in builder.py, which correctly
    handles the RTL jc flip (jc=left = physically RIGHT when bidi=1).
    """
    pPr = paragraph._element.get_or_add_pPr()

    for existing in pPr.findall(qn("w:bidi")):
        pPr.remove(existing)

    if rtl:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        # bidi must come before jc in schema order
        pPr.insert(0, bidi)


def set_paragraph_spacing(
    paragraph: "Paragraph",
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
    line_spacing_pt: float | None = None,
) -> None:
    """Set paragraph spacing properties."""
    fmt = paragraph.paragraph_format
    if space_before_pt is not None:
        fmt.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        fmt.space_after = Pt(space_after_pt)
    if line_spacing_pt is not None:
        from docx.shared import Pt as _Pt
        fmt.line_spacing = _Pt(line_spacing_pt)


def add_paragraph_bottom_border(
    paragraph: "Paragraph",
    color_hex: str = "003366",
    size_eighths_pt: int = 12,   # 12 eighths = 1.5pt
    space: int = 4,
) -> None:
    """Add a bottom border to a paragraph (used for heading separators)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    # Remove old bottom if present
    for old in pBdr.findall(qn("w:bottom")):
        pBdr.remove(old)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths_pt))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bottom)


def add_paragraph_top_border(
    paragraph: "Paragraph",
    color_hex: str = "003366",
    size_eighths_pt: int = 12,
    space: int = 4,
) -> None:
    """Add a top border to a paragraph (used for footer separators)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for old in pBdr.findall(qn("w:top")):
        pBdr.remove(old)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size_eighths_pt))
    top.set(qn("w:space"), str(space))
    top.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(top)


# ---------------------------------------------------------------------------
# Run-level helpers
# ---------------------------------------------------------------------------

def apply_run_formatting(
    run: "Run",
    font_name: str | None = None,
    font_size_pt: int | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
    color_hex: str | None = None,
    rtl: bool = False,
) -> None:
    """Apply formatting to a run, including complex-script (Arabic) font support."""
    if font_name is not None:
        run.font.name = font_name
        # Set complex-script font (w:cs) separately — python-docx only sets ASCII font
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:cs"), font_name)

    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
        # Also set szCs for complex script size
        rPr = run._r.get_or_add_rPr()
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(font_size_pt * 2))  # half-points

    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if underline is not None:
        run.font.underline = underline

    if color_hex is not None:
        from .config import hex_to_rgb
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)

    if rtl:
        _set_run_rtl(run)
        # Re-apply cs font after setting rtl (python-docx bug: rtl can clear w:cs)
        if font_name is not None:
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:cs"), font_name)


def _set_run_rtl(run: "Run") -> None:
    """Add <w:rtl/> to a run's rPr."""
    rPr = run._r.get_or_add_rPr()
    if rPr.find(qn("w:rtl")) is None:
        rtl_el = OxmlElement("w:rtl")
        rPr.append(rtl_el)


# ---------------------------------------------------------------------------
# Cell-level helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell: "_Cell", hex_color: str) -> None:
    """Set cell background color via <w:shd>."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)


def set_cell_borders(
    cell: "_Cell",
    color_hex: str = "BFBFBF",
    width_pt: float = 0.5,
    sides: set[str] | None = None,
) -> None:
    """Set cell borders.  sides: subset of {'top','bottom','left','right','insideH','insideV'}."""
    if sides is None:
        sides = {"top", "bottom", "left", "right"}

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # Size in eighths of a point
    sz = max(1, round(width_pt * 8))
    color = color_hex.lstrip("#").upper()

    for side in sides:
        # Remove existing
        for existing in tcBorders.findall(qn(f"w:{side}")):
            tcBorders.remove(existing)
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)


def set_cell_padding(
    cell: "_Cell",
    top_pt: float = 3.0,
    bottom_pt: float = 3.0,
    left_pt: float = 5.4,
    right_pt: float = 5.4,
) -> None:
    """Set cell inner margins via <w:tcMar>."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    def _twips(pt: float) -> int:
        return round(pt * 20)  # 1 pt = 20 twips

    for side, val_pt in [("top", top_pt), ("bottom", bottom_pt),
                          ("left", left_pt), ("right", right_pt)]:
        for existing in tcMar.findall(qn(f"w:{side}")):
            tcMar.remove(existing)
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(_twips(val_pt)))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)


def fix_table_layout(table) -> None:
    """Force fixed table layout so explicit column widths are respected."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


# ---------------------------------------------------------------------------
# Page number field  (fldChar triplet)
# ---------------------------------------------------------------------------

def add_page_number_field(paragraph: "Paragraph", instr: str = " PAGE ") -> None:
    """Emit the three-run fldChar structure for a PAGE or NUMPAGES field."""
    run = paragraph.add_run()
    # begin
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar_end)


def add_page_x_of_y_field(paragraph: "Paragraph", prefix: str = "Page ") -> None:
    """Emit 'Page X of Y' using two fldChar triplets."""
    if prefix:
        paragraph.add_run(prefix)
    add_page_number_field(paragraph, " PAGE ")
    paragraph.add_run(" of ")
    add_page_number_field(paragraph, " NUMPAGES ")


def add_text_and_page_number_field(
    paragraph: "Paragraph",
    prefix: str = "",
    suffix: str = "",
    font_name: str | None = None,
    font_size_pt: int | None = None,
    rtl: bool = False,
) -> None:
    """Emit a static text prefix, a live PAGE field, and an optional suffix.

    Produces output like ``"Internal use only  |  Page 3"`` where the number
    is a dynamic Word field that updates automatically.

    When ``rtl=True`` two U+200F RIGHT-TO-LEFT MARKs bracket the PAGE field
    (one before, one after).  This anchors the isolated numeral on both sides
    so the Unicode bidi algorithm does not float the LTR digit to the left edge
    of an RTL paragraph.

    Usage via ``DocumentConfig``::

        DocumentConfig(
            footer_type="text_and_page_number",
            footer_text="للاستخدام الرسمي الداخلي  |  صفحة ",
        )
    """
    def _styled_run(text: str) -> None:
        run = paragraph.add_run(text)
        if font_name or font_size_pt:
            apply_run_formatting(run, font_name=font_name, font_size_pt=font_size_pt)

    if prefix:
        _styled_run(prefix)
    if rtl:
        # U+200F before the field: seed the bidi algorithm so the number is
        # treated as part of the RTL sequence, not a free-floating LTR run.
        _styled_run("\u200f")
    add_page_number_field(paragraph, " PAGE ")
    if rtl:
        # U+200F after the field: close the RTL anchor on the trailing side.
        _styled_run("\u200f")
    if suffix:
        _styled_run(suffix)


# ---------------------------------------------------------------------------
# Hyperlink helper
# ---------------------------------------------------------------------------

def add_hyperlink(paragraph: "Paragraph", url: str, display_text: str,
                  font_name: str | None = None, font_size_pt: int | None = None,
                  rtl: bool = False) -> None:
    """Add a clickable hyperlink run to an existing paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style: blue, underline
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "0563C1")
    rPr.append(color_el)

    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)

    if font_name:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:cs"), font_name)
        rPr.insert(0, rFonts)

    if font_size_pt:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(font_size_pt * 2))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(font_size_pt * 2))
        rPr.append(szCs)

    if rtl:
        rtl_el = OxmlElement("w:rtl")
        rPr.append(rtl_el)

    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = display_text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Arabic text helpers
# ---------------------------------------------------------------------------

_ARABIC_RE = re.compile(r"[\u0600-\u06FF]")


def is_arabic(text: str) -> bool:
    return bool(_ARABIC_RE.search(text))


def maybe_reshape(text: str, apply_reshaping: bool) -> str:
    """Optionally apply arabic-reshaper for ligature joining.

    Only call this when apply_arabic_reshaping=True in DocumentConfig.
    Do NOT call get_display() — Word handles bidi reordering natively.
    """
    if not apply_reshaping or not is_arabic(text):
        return text
    try:
        import arabic_reshaper
        return arabic_reshaper.reshape(text)
    except ImportError:
        return text


# ---------------------------------------------------------------------------
# Arabic-Indic numeral conversion
# ---------------------------------------------------------------------------

_ARABIC_INDIC = str.maketrans("0123456789", "٠١٢٣٤٥٦٧٨٩")


def to_arabic_indic(number: int) -> str:
    return str(number).translate(_ARABIC_INDIC)
