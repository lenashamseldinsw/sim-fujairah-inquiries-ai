"""
CoverPage canvas for sword-word-builder.

A CoverPage is a composable blank canvas that accepts the same content
components as WordBuilder. All operations are deferred and replayed on
the live builder when passed to WordBuilder.add_cover_page().
"""
from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Callable, Literal

if TYPE_CHECKING:
    from .builder import WordBuilder
    from .config import TextStyle, TableStyle, ChartStyle


class CoverPage:
    """
    A blank canvas for building a cover page from any combination of components.

    CoverPage mirrors the full WordBuilder content API — any component that can
    be added to a document body (headings, paragraphs, tables, charts, images,
    separators, markdown, etc.) can be placed on a cover page in any order.

    All operations are deferred and executed on the live ``WordBuilder`` instance
    when ``builder.add_cover_page(cover)`` is called. This means charts, tables,
    and all stateful operations work exactly as they do in the main document.

    Usage::

        cover = CoverPage()
        cover.add_spacer(5)
        cover.add_heading("My Report", level=1)
        cover.add_horizontal_separator(native=True, color="2E74B5")
        cover.add_paragraph(
            "Subtitle line",
            style=TextStyle(size=16, italic=True, alignment="CENTER"),
        )
        cover.add_spacer(3)
        cover.add_picture("logo.png", width_cm=6)
        cover.add_table(headers, rows, style=TableStyle(...))
        cover.add_chart("column", data, style=ChartStyle(...))

        builder.add_cover_page(cover)

    For backwards compatibility with pre-2.0 code, use the preset classmethod::

        cover = CoverPage.preset(
            "Report Title",
            subtitle="Subtitle",
            metadata={"Author": "John", "Date": "2025"},
        )
        builder.add_cover_page(cover)
    """

    def __init__(self) -> None:
        self._ops: list[Callable] = []

    # ------------------------------------------------------------------
    # Content methods — mirror WordBuilder's content API
    # ------------------------------------------------------------------

    def add_spacer(self, height_pt: float = 12.0) -> "CoverPage":
        """Insert an empty paragraph with the given space_after (in pt)."""
        self._ops.append(lambda b: b.add_spacer(height_pt))
        return self

    def add_heading(
        self,
        text: str,
        level: int = 1,
        style: "TextStyle | None" = None,
        rtl: bool | None = None,
        separator: bool | None = None,
    ) -> "CoverPage":
        """Add a heading at the given level (1–3)."""
        self._ops.append(
            lambda b: b.add_heading(
                text, level=level, style=style, rtl=rtl, separator=separator
            )
        )
        return self

    def add_paragraph(
        self,
        text: str,
        style: "TextStyle | None" = None,
        rtl: bool | None = None,
        bold: bool = False,
        italic: bool = False,
        alignment: str | None = None,
        space_after: float | None = None,
        keep_with_next: bool = False,
    ) -> "CoverPage":
        """Add a styled paragraph."""
        self._ops.append(
            lambda b: b.add_paragraph(
                text,
                style=style,
                rtl=rtl,
                bold=bold,
                italic=italic,
                alignment=alignment,
                space_after=space_after,
                keep_with_next=keep_with_next,
            )
        )
        return self

    def add_markdown(
        self,
        markdown_text: str,
        rtl: bool | None = None,
        base_style: "TextStyle | None" = None,
    ) -> "CoverPage":
        """Render markdown text."""
        self._ops.append(
            lambda b: b.add_markdown(markdown_text, rtl=rtl, base_style=base_style)
        )
        return self

    def add_hyperlink(
        self,
        url: str,
        display_text: str,
        rtl: bool | None = None,
    ) -> "CoverPage":
        """Add a hyperlink in its own paragraph."""
        self._ops.append(lambda b: b.add_hyperlink(url, display_text, rtl=rtl))
        return self

    def add_picture(
        self,
        image_path: str | Path,
        width_cm: float | None = None,
        height_cm: float | None = None,
        alignment: str = "CENTER",
        caption: str | None = None,
    ) -> "CoverPage":
        """Insert an image."""
        self._ops.append(
            lambda b: b.add_picture(
                image_path,
                width_cm=width_cm,
                height_cm=height_cm,
                alignment=alignment,
                caption=caption,
            )
        )
        return self

    def add_table(
        self,
        data: "list[dict] | list[list] | list[tuple]",
        headers: "list[str] | None" = None,
        style: "TableStyle | None" = None,
        rtl: bool | None = None,
        caption: str | None = None,
        table_type: Literal["data", "metrics"] = "data",
    ) -> "CoverPage":
        """Add a table."""
        self._ops.append(
            lambda b: b.add_table(
                data,
                headers=headers,
                style=style,
                rtl=rtl,
                caption=caption,
                table_type=table_type,
            )
        )
        return self

    def add_chart(
        self,
        data: dict,
        chart_type: Literal["column", "bar", "line", "pie"] = "column",
        style: "ChartStyle | None" = None,
        rtl: bool | None = None,
    ) -> "CoverPage":
        """Add a native Word chart (editable in Word with embedded Excel data).

        data format::

            {
                "title": "My Chart",
                "categories": ["Q1", "Q2", "Q3"],
                "series": [
                    {"name": "Revenue", "values": [100, 150, 130], "color": "4472C4"},
                ],
            }
        """
        self._ops.append(
            lambda b: b.add_chart(data, chart_type=chart_type, style=style, rtl=rtl)
        )
        return self

    def add_horizontal_separator(
        self,
        color: str | None = None,
        native: bool = False,
    ) -> "CoverPage":
        """Add a horizontal separator line.

        Args:
            color: Hex color string (with or without '#'). Defaults to body_color.
            native: When True, uses a real Word paragraph border instead of Unicode
                    characters. Native borders scale correctly and look cleaner in
                    print and PDF export.
        """
        self._ops.append(
            lambda b: b.add_horizontal_separator(color=color, native=native)
        )
        return self

    def add_banner(
        self,
        text: str,
        bg_color: str = "2E74B5",
        text_color: str = "FFFFFF",
        font_size: int = 14,
        bold: bool = True,
        rtl: "bool | None" = None,
    ) -> "CoverPage":
        """Add a full-width single-row coloured banner (section header).

        See :meth:`WordBuilder.add_banner` for full documentation.
        """
        self._ops.append(
            lambda b: b.add_banner(
                text,
                bg_color=bg_color,
                text_color=text_color,
                font_size=font_size,
                bold=bold,
                rtl=rtl,
            )
        )
        return self

    def add_page_break(self) -> "CoverPage":
        """Insert a page break.

        Useful when composing a multi-page cover section before the main body.
        Note that :meth:`WordBuilder.add_cover_page` already appends one
        page break automatically after the cover renders; use this only when
        you need an *additional* break within the cover canvas itself.
        """
        self._ops.append(lambda b: b.add_page_break())
        return self

    # ------------------------------------------------------------------
    # Rendering
    # ------------------------------------------------------------------

    def _render(self, builder: "WordBuilder") -> None:
        """Replay all stored operations on the given builder instance."""
        for op in self._ops:
            op(builder)

    # ------------------------------------------------------------------
    # Preset: backwards-compatible fixed layout
    # ------------------------------------------------------------------

    @classmethod
    def preset(
        cls,
        title: str,
        subtitle: str | None = None,
        metadata: dict[str, str] | None = None,
    ) -> "CoverPage":
        """Create a CoverPage with the classic fixed layout.

        Produces the same output as the old ``builder.add_cover_page(title, ...)``
        API: 5 spacers, large centred title, decorative separator line, optional
        subtitle, 3 spacers, and optional metadata key-value rows.

        Use this to migrate existing code::

            # Before:
            builder.add_cover_page(
                "Report Title",
                subtitle="Subtitle",
                metadata={"Author": "Me"},
            )

            # After:
            cover = CoverPage.preset(
                "Report Title",
                subtitle="Subtitle",
                metadata={"Author": "Me"},
            )
            builder.add_cover_page(cover)
        """
        cover = cls()
        cover._ops.append(lambda b: _render_preset(b, title, subtitle, metadata))
        return cover


# ---------------------------------------------------------------------------
# Preset renderer (module-private)
# ---------------------------------------------------------------------------

def _render_preset(
    builder: "WordBuilder",
    title: str,
    subtitle: str | None,
    metadata: "dict[str, str] | None",
) -> None:
    """Replicate the original fixed cover page layout on a builder instance.

    This is called at render time so it has access to the builder's live config.
    The page break itself is added by WordBuilder.add_cover_page after _render().
    """
    from ._xml_utils import apply_run_formatting, set_paragraph_spacing, is_arabic
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cfg = builder._config
    doc = builder._doc

    # 5 spacers
    for _ in range(5):
        doc.add_paragraph()

    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    apply_run_formatting(
        title_run,
        font_name=cfg.heading_font or cfg.default_font,
        font_size_pt=32,
        bold=True,
        color_hex=cfg.heading_color,
        rtl=is_arabic(title),
    )
    set_paragraph_spacing(title_para, space_before_pt=0, space_after_pt=12)

    # Decorative line
    dec_para = doc.add_paragraph()
    dec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dec_run = dec_para.add_run("━━━━━━━━━━ ◆ ━━━━━━━━━━")
    apply_run_formatting(
        dec_run,
        font_name=cfg.heading_font or cfg.default_font,
        font_size_pt=14,
        color_hex=cfg.accent_color,
    )
    set_paragraph_spacing(dec_para, space_before_pt=12, space_after_pt=24)

    # Subtitle
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitle)
        apply_run_formatting(
            sub_run,
            font_name=cfg.heading_font or cfg.default_font,
            font_size_pt=16,
            italic=True,
            color_hex=cfg.secondary_color,
            rtl=is_arabic(subtitle),
        )
        set_paragraph_spacing(sub_para, space_before_pt=0, space_after_pt=24)

    # 3 spacers
    for _ in range(3):
        doc.add_paragraph()

    # Metadata lines
    if metadata:
        for key, val in metadata.items():
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_text = f"{key}: {val}"
            meta_run = meta_para.add_run(meta_text)
            apply_run_formatting(
                meta_run,
                font_name=cfg.default_font,
                font_size_pt=cfg.default_font_size,
                color_hex=cfg.body_color,
                rtl=is_arabic(meta_text),
            )
            set_paragraph_spacing(meta_para, space_before_pt=0, space_after_pt=4)
