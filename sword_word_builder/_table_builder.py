"""
Internal table construction helpers for sword-word-builder.
"""
from __future__ import annotations

from typing import Any

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import DocumentConfig, TableStyle, CellStyle, CellLine, PAGE_SIZES, hex_to_rgb
from ._xml_utils import (
    set_cell_shading,
    set_cell_borders,
    set_cell_padding,
    fix_table_layout,
    set_paragraph_rtl,
    apply_run_formatting,
    alignment_const,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _usable_width_emu(config: DocumentConfig) -> int:
    """Return the usable page width in EMU (page width minus left+right margins)."""
    page_w, _ = PAGE_SIZES[config.page_size]
    margin_l = int(config.margin_left * 360000)
    margin_r = int(config.margin_right * 360000)
    return int(page_w) - margin_l - margin_r


def _unwrap_cell(value: Any) -> tuple[Any, CellStyle | None]:
    """Return (raw_value, CellStyle | None) by unwrapping (value, CellStyle) tuples.

    If value is a list[CellLine], it is returned as-is with no CellStyle.
    """
    if isinstance(value, list) and value and isinstance(value[0], CellLine):
        return value, None
    if isinstance(value, tuple) and len(value) == 2 and isinstance(value[1], CellStyle):
        return value[0], value[1]
    return value, None


def _normalize_data(
    data: list[dict] | list[list] | list[tuple],
    headers: list[str] | None,
    table_type: str,
) -> tuple[list[str], list[list[Any]]]:
    """Return (headers, rows) as plain lists regardless of input format.

    Each element in the returned rows may be either a plain value or a
    ``(value, CellStyle)`` tuple — callers should use ``_unwrap_cell`` when
    consuming individual cells.
    """
    if table_type == "metrics":
        # data is list[tuple[label, value]] or dict
        if isinstance(data, dict):
            items = list(data.items())
        else:
            items = [(str(k), v) for k, v in data]
        # Honour caller-supplied header labels; fall back to generic English names.
        col_headers = headers if headers is not None else ["Label", "Value"]
        return col_headers, [[k, v] for k, v in items]

    if not data:
        return headers or [], []

    first = data[0]
    if isinstance(first, dict):
        hdrs = headers if headers is not None else list(first.keys())
        rows = [[row.get(h, "") for h in hdrs] for row in data]
        return hdrs, rows
    else:
        # list[list] or list[tuple]
        if headers is None:
            raise ValueError(
                "headers= is required when data is a list of lists/tuples."
            )
        return list(headers), [list(r) for r in data]


def build_table(
    doc: Document,
    data: list[dict] | list[list] | list[tuple],
    headers: list[str] | None,
    style: TableStyle,
    config: DocumentConfig,
    rtl: bool,
    table_type: str = "data",
) -> Any:  # docx.table.Table
    """Build and insert a table into the document.  Returns the Table object."""
    hdrs, rows = _normalize_data(data, headers, table_type)
    n_cols = len(hdrs)
    if n_cols == 0:
        return None
    n_rows = len(rows)

    # --- Create table ---
    table = doc.add_table(rows=0, cols=n_cols)

    # Use Table Grid as base style (universally available)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    fix_table_layout(table)

    # --- Table horizontal alignment ---
    if style.table_alignment != "LEFT":
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)
        for old_jc in tblPr.findall(qn("w:jc")):
            tblPr.remove(old_jc)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), style.table_alignment.lower())
        tblPr.append(jc)

    # --- Column widths ---
    usable = _usable_width_emu(config)
    if style.column_widths and len(style.column_widths) == n_cols:
        widths = [int(f * usable) for f in style.column_widths]
    else:
        equal = usable // n_cols
        widths = [equal] * n_cols

    # --- Header row ---
    hdr_row = table.add_row()
    for col_idx, hdr_text_raw in enumerate(hdrs):
        cell = hdr_row.cells[col_idx]
        cell.width = widths[col_idx]
        # Per-column bg override takes precedence over the single header_bg_color.
        hdr_bg = (
            style.header_bg_colors[col_idx]
            if style.header_bg_colors and col_idx < len(style.header_bg_colors)
            else style.header_bg_color
        )
        hdr_text, _ = _unwrap_cell(hdr_text_raw)
        if isinstance(hdr_text, list) and hdr_text and isinstance(hdr_text[0], CellLine):
            _format_cell_multiline(
                cell, hdr_text,
                bg_color=hdr_bg,
                font_name=config.default_font,
                rtl=rtl,
                style=style,
            )
        else:
            _format_cell(
                cell, str(hdr_text),
                bg_color=hdr_bg,
                text_color=style.header_text_color,
                bold=style.header_bold,
                font_size=style.header_font_size or config.default_font_size,
                font_name=config.default_font,
                alignment=style.text_alignment,
                rtl=rtl,
                style=style,
                is_header=True,
            )
        _apply_borders(cell, style, is_perimeter=(True, col_idx == 0, True, col_idx == n_cols - 1))

    # --- Data rows ---
    for row_idx, row_data in enumerate(rows):
        is_alt = (row_idx % 2 == 1) and (style.alt_row_bg_color is not None)
        row_bg = style.alt_row_bg_color if is_alt else style.row_bg_color

        tbl_row = table.add_row()
        for col_idx, cell_val_raw in enumerate(row_data):
            if col_idx >= n_cols:
                break
            cell = tbl_row.cells[col_idx]
            cell.width = widths[col_idx]

            # Unwrap (value, CellStyle) tuples
            cell_val, cell_override = _unwrap_cell(cell_val_raw)

            # Resolve background — CellStyle.bg_color > row bg
            bg = cell_override.bg_color if (cell_override and cell_override.bg_color) else row_bg

            # Resolve text color — CellStyle.text_color > TableStyle default
            text_color = (
                cell_override.text_color
                if (cell_override and cell_override.text_color)
                else style.cell_text_color
            )

            # Resolve font size — CellStyle.font_size > TableStyle default
            font_size = (
                cell_override.font_size
                if (cell_override and cell_override.font_size is not None)
                else (style.font_size or config.default_font_size)
            )

            # Auto right-align numbers
            if style.auto_align_numbers and isinstance(cell_val, (int, float)):
                cell_align = "RIGHT"
            else:
                cell_align = style.text_alignment

            # Metrics table: label column alignment; bold uses TableStyle.metrics_label_bold
            if table_type == "metrics" and col_idx == 0:
                cell_bold = style.metrics_label_bold
                cell_align = "RIGHT"
            elif table_type == "metrics" and col_idx == 1:
                cell_bold = False
                cell_align = "LEFT"
            else:
                cell_bold = False

            # CellStyle.bold overrides everything above
            if cell_override and cell_override.bold is not None:
                cell_bold = cell_override.bold

            if isinstance(cell_val, list) and cell_val and isinstance(cell_val[0], CellLine):
                _format_cell_multiline(
                    cell, cell_val,
                    bg_color=bg,
                    font_name=config.default_font,
                    rtl=rtl,
                    style=style,
                )
            else:
                _format_cell(
                    cell, str(cell_val) if cell_val is not None else "",
                    bg_color=bg,
                    text_color=text_color,
                    bold=cell_bold,
                    font_size=font_size,
                    font_name=config.default_font,
                    alignment=cell_align,
                    rtl=rtl,
                    style=style,
                    is_header=False,
                )
            is_top = row_idx == 0
            is_bottom = row_idx == n_rows - 1
            is_left = col_idx == 0
            is_right = col_idx == n_cols - 1
            _apply_borders(cell, style, is_perimeter=(is_top, is_left, is_bottom, is_right))

    return table


def _format_cell(
    cell,
    text: str,
    bg_color: str,
    text_color: str,
    bold: bool,
    font_size: int,
    font_name: str,
    alignment: str,
    rtl: bool,
    style: TableStyle,
    is_header: bool,
) -> None:
    # Clear default empty paragraph
    para = cell.paragraphs[0]
    para.clear()

    if rtl:
        set_paragraph_rtl(para, True)
    else:
        para.alignment = alignment_const(alignment)

    set_cell_shading(cell, bg_color)
    set_cell_padding(cell,
                     top_pt=style.cell_padding_top,
                     bottom_pt=style.cell_padding_bottom,
                     left_pt=style.cell_padding_left,
                     right_pt=style.cell_padding_right)

    run = para.add_run(text)
    apply_run_formatting(
        run,
        font_name=font_name,
        font_size_pt=font_size,
        bold=bold,
        color_hex=text_color,
        rtl=rtl,
    )

    from .config import DocumentConfig
    from docx.shared import Pt
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)


def _format_cell_multiline(
    cell,
    lines: list[CellLine],
    bg_color: str,
    font_name: str,
    rtl: bool,
    style: TableStyle,
) -> None:
    """Render a list[CellLine] as multiple <w:p> elements inside a single <w:tc>."""
    set_cell_shading(cell, bg_color)
    set_cell_padding(cell,
                     top_pt=style.cell_padding_top,
                     bottom_pt=style.cell_padding_bottom,
                     left_pt=style.cell_padding_left,
                     right_pt=style.cell_padding_right)

    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
            para.clear()
        else:
            para = cell.add_paragraph()

        # Tight spacing between stacked lines
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        # Alignment — set RTL flag AND explicit alignment independently
        eff_align = line.alignment or style.text_alignment
        eff_rtl = line.rtl if line.rtl is not None else rtl
        if eff_rtl:
            set_paragraph_rtl(para, True)
        para.alignment = alignment_const(eff_align)

        # Effective run values (line overrides style defaults)
        eff_font_size = line.font_size or style.font_size
        eff_bold = line.bold if line.bold is not None else False
        eff_italic = line.italic if line.italic is not None else False
        eff_color = line.color or style.cell_text_color

        run = para.add_run(line.text)
        apply_run_formatting(
            run,
            font_name=font_name,
            font_size_pt=eff_font_size,
            bold=eff_bold,
            italic=eff_italic,
            color_hex=eff_color,
            rtl=eff_rtl,
        )


def _apply_borders(cell, style: TableStyle, is_perimeter: tuple[bool, bool, bool, bool]) -> None:
    """Apply borders to a cell.  is_perimeter = (top, left, bottom, right)."""
    is_top, is_left, is_bottom, is_right = is_perimeter

    sides: set[str] = set()
    if style.show_outer_border:
        if is_top:
            sides.add("top")
        if is_left:
            sides.add("left")
        if is_bottom:
            sides.add("bottom")
        if is_right:
            sides.add("right")

    if style.show_inner_borders:
        if not is_top:
            sides.add("top")
        if not is_left:
            sides.add("left")
        if not is_bottom:
            sides.add("bottom")
        if not is_right:
            sides.add("right")

    if sides:
        set_cell_borders(cell, color_hex=style.border_color,
                         width_pt=style.border_width_pt, sides=sides)
