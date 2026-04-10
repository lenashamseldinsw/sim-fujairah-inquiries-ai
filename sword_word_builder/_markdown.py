"""
Markdown-to-docx renderer for sword-word-builder.

Supports:
  Block:  # headings (1–3), - bullet lists, 1. numbered lists, plain paragraphs
  Inline: **bold**, *italic*, ***bold+italic***, [text](url), bare https?:// URLs
  RTL:    Arabic-Indic numerals in numbered lists, auto-bold specific Arabic phrases,
          w:bidi + alignment set on every RTL paragraph
"""
from __future__ import annotations

import re
from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ._xml_utils import (
    apply_run_formatting,
    set_paragraph_rtl,
    add_hyperlink,
    is_arabic,
    maybe_reshape,
)

if TYPE_CHECKING:
    from docx import Document
    from .config import DocumentConfig, TextStyle


# ---------------------------------------------------------------------------
# Inline token patterns
# ---------------------------------------------------------------------------

# Order matters: process longest/most-specific patterns first
_INLINE_PATTERNS = [
    ("link",     re.compile(r"\[([^\]]+)\]\((https?://[^\)]+)\)")),
    ("bold_italic", re.compile(r"\*{3}(.+?)\*{3}")),
    ("bold",     re.compile(r"\*{2}(.+?)\*{2}")),
    ("italic",   re.compile(r"\*(.+?)\*")),
    ("url",      re.compile(r"(https?://\S+)")),
]

# Numbered list: optional leading **, digit(s) or Arabic-Indic, dot, optional **, space, text
_NUM_LIST_RE = re.compile(r"^(\*\*)?([٠-٩\d]+)\.(\*\*)?\s+(.+)")
# Bullet list
_BULLET_RE = re.compile(r"^[•\-\*]\s+(.+)")
# Heading
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)")
# Continuation line (indented, belongs to previous list item)
_CONTINUATION_RE = re.compile(r"^\s{2,}(.+)")


class MarkdownRenderer:
    """Renders a markdown string into an open python-docx Document."""

    def __init__(self, config: "DocumentConfig", base_style: "TextStyle | None" = None):
        self._config = config
        self._base_style = base_style

    def render(self, doc: "Document", text: str, rtl: bool) -> None:
        """Parse and render markdown text into the document."""
        lines = text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]

            # --- Heading ---
            m = _HEADING_RE.match(line)
            if m:
                level = min(len(m.group(1)), 3)
                self._render_heading(doc, m.group(2).strip(), level, rtl)
                i += 1
                continue

            # --- Bullet ---
            m = _BULLET_RE.match(line)
            if m:
                self._render_bullet(doc, m.group(1).strip(), rtl)
                i += 1
                continue

            # --- Numbered list ---
            m = _NUM_LIST_RE.match(line)
            if m:
                num_str = m.group(2)
                content = m.group(4).strip()
                # Gather continuation lines
                i += 1
                while i < len(lines) and _CONTINUATION_RE.match(lines[i]):
                    content += " " + lines[i].strip()
                    i += 1
                number = int("".join(
                    "0123456789"[ord(c) - ord("٠")] if "٠" <= c <= "٩" else c
                    for c in num_str
                ))
                self._render_numbered(doc, number, content, rtl)
                continue

            # --- Blank line → spacer ---
            if line.strip() == "":
                i += 1
                continue

            # --- Paragraph ---
            self._render_paragraph(doc, line.strip(), rtl)
            i += 1

    # ------------------------------------------------------------------
    # Block renderers
    # ------------------------------------------------------------------

    def _render_heading(self, doc: "Document", text: str, level: int, rtl: bool) -> None:
        cfg = self._config
        size_map = {1: cfg.heading1_size, 2: cfg.heading2_size, 3: cfg.heading3_size}
        size = size_map.get(level, cfg.default_font_size + 2)
        font = cfg.heading_font or cfg.default_font

        para = doc.add_paragraph(style=f"Heading {level}")
        para.clear()
        if rtl:
            set_paragraph_rtl(para, True)
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = para.add_run(_reshape(text, cfg))
        apply_run_formatting(
            run,
            font_name=font,
            font_size_pt=size,
            bold=cfg.heading_bold,
            color_hex=cfg.accent_color,
            rtl=rtl,
        )
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(4)

    def _render_bullet(self, doc: "Document", text: str, rtl: bool) -> None:
        try:
            para = doc.add_paragraph(style="List Bullet")
        except Exception:
            para = doc.add_paragraph()

        para.clear()
        if rtl:
            set_paragraph_rtl(para, True)
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        self._add_formatted_runs(para, text, rtl)
        para.paragraph_format.space_after = Pt(2)

    def _render_numbered(self, doc: "Document", number: int, text: str, rtl: bool) -> None:
        para = doc.add_paragraph()
        if rtl:
            set_paragraph_rtl(para, True)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        num_display = str(number)  # always use Western Arabic numerals (1, 2, 3)

        # Remove indentation
        fmt = para.paragraph_format
        fmt.left_indent = Pt(0)
        fmt.right_indent = Pt(0)
        fmt.first_line_indent = Pt(0)
        fmt.space_after = Pt(6)

        cfg = self._config
        num_run = para.add_run(f"{num_display}. ")
        apply_run_formatting(
            num_run,
            font_name=cfg.default_font,
            font_size_pt=cfg.default_font_size,
            bold=True,
            rtl=rtl,
        )

        self._add_formatted_runs(para, text, rtl)

    def _render_paragraph(self, doc: "Document", text: str, rtl: bool) -> None:
        cfg = self._config
        text = self._auto_bold_arabic(text) if rtl else text

        para = doc.add_paragraph()
        if rtl:
            set_paragraph_rtl(para, True)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        para.paragraph_format.space_after = Pt(cfg.default_paragraph_space_after)
        self._add_formatted_runs(para, text, rtl)

    # ------------------------------------------------------------------
    # Inline formatting
    # ------------------------------------------------------------------

    def _add_formatted_runs(self, para, text: str, rtl: bool) -> None:
        """Parse inline markdown and add appropriately formatted runs."""
        cfg = self._config
        remaining = text

        while remaining:
            earliest_match = None
            earliest_start = len(remaining)
            matched_type = None

            for token_type, pattern in _INLINE_PATTERNS:
                m = pattern.search(remaining)
                if m and m.start() < earliest_start:
                    earliest_match = m
                    earliest_start = m.start()
                    matched_type = token_type

            # Emit text before the match as a plain run
            if earliest_start > 0:
                plain = remaining[:earliest_start]
                plain = _reshape(plain, cfg)
                run = para.add_run(plain)
                apply_run_formatting(
                    run,
                    font_name=cfg.default_font,
                    font_size_pt=cfg.default_font_size,
                    color_hex=cfg.body_color,
                    rtl=rtl,
                )

            if earliest_match is None:
                break

            m = earliest_match

            if matched_type == "link":
                display = m.group(1)
                url = m.group(2)
                is_rtl_link = rtl and is_arabic(display)
                add_hyperlink(para, url, display,
                              font_name=cfg.default_font,
                              font_size_pt=cfg.default_font_size,
                              rtl=is_rtl_link)

            elif matched_type == "url":
                url = m.group(1)
                add_hyperlink(para, url, url,
                              font_name=cfg.default_font,
                              font_size_pt=cfg.default_font_size,
                              rtl=False)  # bare URLs stay LTR

            elif matched_type == "bold_italic":
                run = para.add_run(_reshape(m.group(1), cfg))
                apply_run_formatting(run, font_name=cfg.default_font,
                                     font_size_pt=cfg.default_font_size,
                                     bold=True, italic=True, rtl=rtl)

            elif matched_type == "bold":
                run = para.add_run(_reshape(m.group(1), cfg))
                apply_run_formatting(run, font_name=cfg.default_font,
                                     font_size_pt=cfg.default_font_size,
                                     bold=True, rtl=rtl)

            elif matched_type == "italic":
                run = para.add_run(_reshape(m.group(1), cfg))
                apply_run_formatting(run, font_name=cfg.default_font,
                                     font_size_pt=cfg.default_font_size,
                                     italic=True, rtl=rtl)

            remaining = remaining[m.end():]

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _auto_bold_arabic(self, text: str) -> str:
        """Wrap configured Arabic phrase starters in **bold** markers."""
        for phrase in self._config.arabic_auto_bold_phrases:
            if text.strip().startswith(phrase):
                text = text.replace(phrase, f"**{phrase}**", 1)
                break
        return text


def _reshape(text: str, config) -> str:
    return maybe_reshape(text, config.apply_arabic_reshaping)
