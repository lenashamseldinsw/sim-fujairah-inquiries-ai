"""Extract tables and charts exactly as they appear in the report.

LEGACY MODULE: This module is maintained for backward compatibility.
New code should use analysis.adaptive_extractor module instead.
"""

import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Any
from chart_parser import extract_charts_from_docx
from analysis.demo.adaptive_extractor import AdaptiveReportExtractor


# Re-export the adaptive extractor function for backward compatibility
def extract_full_report(docx_path: str) -> Dict[str, Any]:
    """
    Extract full report with automatic structure detection and caching.
    
    This function now uses the new adaptive extraction system while maintaining
    backward compatibility with the old API.
    
    Args:
        docx_path: Path to Word document
        
    Returns:
        Dictionary with complete report structure
    """
    extractor = AdaptiveReportExtractor()
    return extractor.extract_report(docx_path)


# ============================================================================
# LEGACY CODE BELOW - Kept for reference and manual extraction if needed
# ============================================================================


def extract_full_report(docx_path: str) -> Dict[str, Any]:
    """Extract all 8 sections with complete table data and native charts."""
    doc = Document(docx_path)

    # Extract native charts from document
    charts = extract_charts_from_docx(docx_path)

    report = {
        'title': '',
        'charts': charts,
        'sections': {
            'executive_summary': {
                'title': 'الملخص التنفيذي',
                'title_en': 'Executive Summary',
                'content': '',
                'key_message': '',
                'tables': []
            },
            'analysis_1': {
                'title': 'التحليل الأول — خريطة عبء العمل الحقيقي',
                'title_en': 'Analysis 1: Workload Distribution',
                'tables': []
            },
            'analysis_2': {
                'title': 'التحليل الثاني — التحديات في رحلة المتعامل',
                'title_en': 'Analysis 2: Customer Journey Challenges',
                'content': '',
                'tables': []
            },
            'analysis_3': {
                'title': 'التحليل الثالث — تحليل الفجوات الرقمية',
                'title_en': 'Analysis 3: Digital Gaps',
                'tables': []
            },
            'analysis_4': {
                'title': 'التحليل الرابع — استراتيجية التحويل الرقمي',
                'title_en': 'Analysis 4: Digital Transformation Strategy',
                'tables': []
            },
            'use_cases': {
                'title': 'حالات الاستخدام المدعومة بالذكاء الاصطناعي',
                'title_en': 'AI-Supported Use Cases',
                'tables': []
            },
            'roadmap': {
                'title': 'خارطة الطريق الاستراتيجية المتكاملة',
                'title_en': 'Strategic Roadmap',
                'tables': []
            },
            'conclusion': {
                'title': 'الخلاصة — من البيانات إلى القرار',
                'title_en': 'Conclusion: From Data to Decision',
                'tables': []
            }
        }
    }

    # Extract title
    for para in doc.paragraphs[:10]:
        if 'تقرير' in para.text and 'استفسارات' in para.text:
            report['title'] = para.text.strip()
            break

    # Extract text
    full_text = '\n'.join([p.text for p in doc.paragraphs])

    # Extract key message
    key_msg_pattern = r'الرسالة الجوهرية:?\s*(.+?)(?=\n\n|\nالنتائج|\nثانياً)'
    match = re.search(key_msg_pattern, full_text, re.DOTALL)
    if match:
        report['sections']['executive_summary']['key_message'] = match.group(1).strip()

    # Extract all tables with exact structure
    extract_all_tables(doc, report)

    return report


def extract_all_tables(doc: Document, report: Dict) -> None:
    """Extract all tables from document with exact column names and data."""

    # Map table index to section and subsection
    # Tables are 0-indexed, but displayed as 1-indexed
    table_mapping = {
        2: ('executive_summary', 'Key Discoveries'),                    # Table 3
        3: ('analysis_1', 'Communication Types'),                       # Table 4
        5: ('analysis_1', 'Misclassified Cases Examples'),             # Table 6
        7: ('analysis_2', 'Friction Points'),                          # Table 8
        9: ('analysis_3', 'Digital Gaps'),                             # Table 10
        10: ('analysis_3', 'Root Causes'),                             # Table 11
        11: ('analysis_4', 'FAQ Questions'),                           # Table 12
        12: ('analysis_4', 'Notification Strategy'),                   # Table 13
        13: ('use_cases', 'AI Tools'),                                 # Table 14
        14: ('roadmap', 'Strategic Recommendations'),                  # Table 15
        16: ('conclusion', 'Transformation Axes'),                     # Table 17
        17: ('conclusion', 'Impact Metrics')                           # Table 18
    }

    for table_idx, table in enumerate(doc.tables):
        if table_idx in table_mapping:
            section_key, table_title = table_mapping[table_idx]

            # Extract table with exact structure
            table_data = extract_table_exact(table)

            if table_data['rows']:
                table_data['title'] = table_title
                report['sections'][section_key]['tables'].append(table_data)


def extract_table_exact(table) -> Dict[str, Any]:
    """Extract table preserving exact column names and structure."""

    # Get column names from first row
    columns = []
    if len(table.rows) > 0:
        first_row = table.rows[0]
        columns = [cell.text.strip() for cell in first_row.cells]

    # Get data rows
    rows = []
    for row in table.rows[1:]:  # Skip header row
        row_data = {}
        cells = [cell.text.strip() for cell in row.cells]

        for col_idx, col_name in enumerate(columns):
            if col_idx < len(cells):
                row_data[col_name] = cells[col_idx]
            else:
                row_data[col_name] = ''

        rows.append(row_data)

    return {
        'columns': columns,
        'rows': rows,
        'title': '',
        'row_count': len(rows)
    }
