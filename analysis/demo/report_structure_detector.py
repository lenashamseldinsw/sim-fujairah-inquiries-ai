"""Automatically detect report structure from Word documents."""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, List, Any, Tuple, Optional
import re


class ReportStructureDetector:
    """
    Automatically detects sections, tables, and charts from Word documents.

    This makes the system adaptable to different report formats without hardcoding.
    """

    # Common heading patterns (Arabic and English)
    HEADING_PATTERNS = [
        r'(التحليل|تحليل|Analysis)',
        r'(الملخص|ملخص|Summary)',
        r'(الخلاصة|خلاصة|Conclusion)',
        r'(التوصيات|توصيات|Recommendations)',
        r'(المقدمة|مقدمة|Introduction)',
        r'(حالات الاستخدام|Use Cases)',
        r'(خارطة الطريق|Roadmap)',
    ]

    def __init__(self, doc_path: str):
        """Initialize detector with document path."""
        self.doc_path = doc_path
        self.doc = Document(doc_path)

    def detect_structure(self) -> Dict[str, Any]:
        """
        Detect complete report structure with hierarchical sections and subsections.

        Returns:
            Dictionary with main sections, subsections, tables, and metadata
        """
        # First detect all headings (main + sub)
        all_headings = self._detect_all_headings()

        # Organize into hierarchical structure
        main_sections = self._organize_hierarchy(all_headings)

        # Detect tables
        tables = self._detect_tables()
        charts = self._detect_chart_positions()

        # Assign tables and charts to sections/subsections
        self._assign_elements_to_hierarchy(main_sections, tables, charts)

        return {
            'sections': main_sections,
            'tables': tables,
            'charts': charts,
            'metadata': self._extract_metadata()
        }

    def _detect_all_headings(self) -> List[Dict[str, Any]]:
        """
        Detect all headings (main sections and subsections).

        Returns:
            List of all headings with their metadata
        """
        headings = []

        for idx, para in enumerate(self.doc.paragraphs):
            heading_type = self._classify_heading(para)

            if heading_type:
                heading = {
                    'id': self._generate_section_id(para.text, len(headings)),
                    'title_ar': para.text.strip(),
                    'title_en': self._translate_section_title(para.text),
                    'position': idx,
                    'type': heading_type,  # 'main' or 'sub'
                    'level': self._get_heading_level(para),
                    'tables': [],
                    'charts': [],
                    'subsections': [],
                    'content': ''
                }
                headings.append(heading)

        return headings

    def _organize_hierarchy(self, headings: List[Dict]) -> List[Dict]:
        """
        Organize flat list of headings into hierarchical structure.

        Main sections contain their subsections, matched by section number.
        Removes duplicate main sections (e.g., from table of contents).

        Returns:
            List of main sections with nested subsections
        """
        # First, collect all main sections and subsections separately
        main_sections_by_title = {}
        main_sections_order = []
        subsections = []

        for heading in headings:
            if heading['type'] == 'main':
                title_key = heading['title_ar'].strip()

                # Only keep first occurrence (skip table of contents duplicates)
                if title_key not in main_sections_by_title:
                    heading['subsections'] = []
                    main_sections_by_title[title_key] = heading
                    main_sections_order.append(heading)

            elif heading['type'] == 'sub':
                subsections.append(heading)

        # Now assign subsections to their parent main section by matching section numbers
        # E.g., "2.1" belongs to "ثانياً", "3.2" belongs to "ثالثاً", etc.
        arabic_ordinals = {
            'أولاً': 1,
            'ثانياً': 2,
            'ثالثاً': 3,
            'رابعاً': 4,
            'خامساً': 5,
            'سادساً': 6,
            'سابعاً': 7,
            'ثامناً': 8,
            'تاسعاً': 9,
            'عاشراً': 10,
        }

        # Create mapping from section number to main section
        main_sections_by_number = {}
        for main_sec in main_sections_order:
            for ordinal, num in arabic_ordinals.items():
                if ordinal in main_sec['title_ar']:
                    main_sections_by_number[num] = main_sec
                    break

        # Assign subsections based on their number prefix (e.g., "2.1" -> section 2)
        for subsec in subsections:
            # Extract section number from subsection title (e.g., "2.1" -> 2)
            match = re.match(r'^(\d+)\.', subsec['title_ar'])
            if match:
                section_num = int(match.group(1))
                parent_section = main_sections_by_number.get(section_num)

                if parent_section:
                    parent_section['subsections'].append(subsec)
                # If no parent found, skip this subsection

        # Remove duplicate main sections - keep the one with subsections
        # Also filter out non-ordinal sections (like document title)
        final_sections = []
        seen_numbers = set()  # Track by Arabic ordinal number

        arabic_ordinals_list = ['أولاً', 'ثانياً', 'ثالثاً', 'رابعاً', 'خامساً',
                                'سادساً', 'سابعاً', 'ثامناً', 'تاسعاً', 'عاشراً']

        for main_sec in main_sections_order:
            # Check if this section has an Arabic ordinal
            ordinal_found = None
            for ordinal in arabic_ordinals_list:
                if ordinal in main_sec['title_ar']:
                    ordinal_found = ordinal
                    break

            # Only keep sections with Arabic ordinals (filters out title page, etc.)
            if ordinal_found:
                if ordinal_found not in seen_numbers:
                    final_sections.append(main_sec)
                    seen_numbers.add(ordinal_found)
                else:
                    # This ordinal already exists - merge if this one has more subsections
                    existing_sec = next(s for s in final_sections if ordinal_found in s['title_ar'])
                    if len(main_sec['subsections']) > len(existing_sec['subsections']):
                        # This duplicate has more subsections, replace
                        idx = final_sections.index(existing_sec)
                        final_sections[idx] = main_sec
                    elif main_sec['subsections']:
                        # Both have subsections, merge them
                        existing_sec['subsections'].extend(main_sec['subsections'])

        # Extract content for each section and subsection
        self._extract_content_for_hierarchy(final_sections)

        return final_sections

    def _extract_content_for_hierarchy(self, main_sections: List[Dict]) -> None:
        """
        Extract content for main sections and subsections.

        Content is text between heading and next heading.
        """
        # Create a flat list of all positions (sections + subsections)
        all_positions = []

        for main_sec in main_sections:
            all_positions.append((main_sec['position'], main_sec))
            for subsec in main_sec['subsections']:
                all_positions.append((subsec['position'], subsec))

        # Sort by position
        all_positions.sort(key=lambda x: x[0])

        # Extract content between positions
        for i, (pos, section) in enumerate(all_positions):
            start_pos = pos + 1
            end_pos = all_positions[i + 1][0] if i + 1 < len(all_positions) else len(self.doc.paragraphs)

            content_paras = []
            for para_idx in range(start_pos, end_pos):
                if para_idx < len(self.doc.paragraphs):
                    para = self.doc.paragraphs[para_idx]
                    text = para.text.strip()
                    if text and not self._classify_heading(para):
                        content_paras.append(text)

            section['content'] = '\n'.join(content_paras)

    def _detect_tables(self) -> List[Dict[str, Any]]:
        """
        Detect all tables with their position in document.

        Returns:
            List of table metadata with position info
        """
        tables_info = []

        # Get paragraph positions for reference
        para_positions = {id(para._element): idx for idx, para in enumerate(self.doc.paragraphs)}

        for table_idx, table in enumerate(self.doc.tables):
            # Find table position by looking at surrounding elements
            table_position = self._find_table_position(table, table_idx)

            # Extract table data
            table_data = self._extract_table_data(table)

            tables_info.append({
                'index': table_idx,
                'position': table_position,
                'data': table_data,
                'assigned_section': None  # Will be assigned later
            })

        return tables_info

    def _detect_chart_positions(self) -> List[Dict[str, Any]]:
        """
        Detect chart positions in document by finding drawingML elements.

        Uses multiple detection methods to find charts/images/shapes:
        - Look for runs containing drawing elements
        - Check for alternative content (AlternateContent)
        - Check for embedded shapes and objects

        Returns:
            List of chart positions with proper positioning
        """
        chart_positions = []
        para_count = 0
        chart_count = 0

        # Traverse document body in actual order
        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                # Check for any inline shapes or drawing elements in this paragraph
                has_chart = self._paragraph_has_visual_content(element)

                if has_chart:
                    chart_positions.append({
                        'position': para_count,
                        'index': chart_count,
                        'assigned_section': None
                    })
                    print(f"  📊 Chart detected at position {para_count} (index {chart_count})")
                    chart_count += 1

                para_count += 1
            elif element.tag.endswith('tbl'):
                # Check tables for embedded content
                if self._element_has_visual_content(element):
                    chart_positions.append({
                        'position': para_count,
                        'index': chart_count,
                        'assigned_section': None
                    })
                    print(f"  📊 Chart detected in table at position {para_count} (index {chart_count})")
                    chart_count += 1

        return chart_positions

    def _paragraph_has_visual_content(self, para_elem) -> bool:
        """Check if paragraph contains any visual content (drawings, images, shapes)."""
        # Check for drawing elements in runs
        for run in para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            # Check for drawingML (charts, shapes, images)
            if run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None:
                return True

        # Check for alternate content (used for compatibility)
        if para_elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}AlternateContent') is not None:
            return True

        # Check for objects
        if para_elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object') is not None:
            return True

        return False

    def _element_has_visual_content(self, elem) -> bool:
        """Check if any element contains visual content."""
        # Check for drawing elements
        if elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None:
            return True
        if elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}AlternateContent') is not None:
            return True
        if elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object') is not None:
            return True
        return False

    def _assign_elements_to_hierarchy(
        self,
        main_sections: List[Dict],
        tables: List[Dict],
        charts: List[Dict]
    ) -> None:
        """
        Assign tables and charts to sections/subsections based on proximity.

        Modifies sections, subsections, tables, and charts in-place.
        """
        if not main_sections:
            return

        # Build a flat list of all section/subsection positions for proximity matching
        position_map = []

        for main_sec in main_sections:
            position_map.append({
                'position': main_sec['position'],
                'section': main_sec,
                'is_main': True
            })

            for subsec in main_sec['subsections']:
                position_map.append({
                    'position': subsec['position'],
                    'section': subsec,
                    'is_main': False,
                    'parent': main_sec
                })

        # Sort by position
        position_map.sort(key=lambda x: x['position'])

        # Assign tables
        for table_info in tables:
            table_pos = table_info['position']

            # Find the closest section/subsection before this table
            assigned_to = None
            for i, pos_info in enumerate(position_map):
                next_pos = position_map[i + 1]['position'] if i + 1 < len(position_map) else float('inf')

                if pos_info['position'] <= table_pos < next_pos:
                    assigned_to = pos_info['section']
                    assigned_to['tables'].append(table_info)
                    table_info['assigned_section'] = assigned_to['id']
                    break

        # Assign charts (usually at document level, add to first section)
        for chart_info in charts:
            chart_pos = chart_info['position']

            assigned_to = None
            for i, pos_info in enumerate(position_map):
                next_pos = position_map[i + 1]['position'] if i + 1 < len(position_map) else float('inf')

                if pos_info['position'] <= chart_pos < next_pos:
                    assigned_to = pos_info['section']
                    assigned_to['charts'].append(chart_info)
                    chart_info['assigned_section'] = assigned_to['id']
                    print(f"  📊 Chart at position {chart_pos} → {assigned_to['title_ar']}")
                    break

            # If no assignment found, add to first main section
            if not assigned_to and main_sections:
                main_sections[0]['charts'].append(chart_info)
                chart_info['assigned_section'] = main_sections[0]['id']
                print(f"  📊 Chart at position {chart_pos} → {main_sections[0]['title_ar']} (fallback)")

    def _classify_heading(self, para) -> str:
        """
        Classify a paragraph as 'main' section, 'sub' section, or None.

        Uses Word heading styles as primary method, with pattern fallbacks.
        This is more reliable than font size which can vary.

        Returns:
            'main' - Main section (Heading 1, or أولاً, ثانياً, etc.)
            'sub' - Subsection (Heading 2, or numbered like 2.1, 2.2)
            None - Not a heading
        """
        text = para.text.strip()
        if not text:
            return None

        # Skip if text is too long (likely paragraph content)
        if len(text) > 300:
            return None

        # Skip table of contents entries (small font in TOC vs larger in actual content)
        # TOC entries are typically 9-10pt, actual headings are 13-16pt
        if para.runs and para.runs[0].font.size:
            font_size = para.runs[0].font.size.pt
            if font_size < 11:  # Skip TOC entries
                return None

        # PRIMARY METHOD: Check Word heading styles (most reliable)
        if para.style and 'Heading' in para.style.name:
            if 'Heading 1' in para.style.name:
                return 'main'
            elif 'Heading 2' in para.style.name or 'Heading 3' in para.style.name:
                return 'sub'

        # FALLBACK: Check for main section patterns (Arabic ordinals with colon)
        # This catches cases where heading styles aren't applied consistently
        main_patterns = [
            r'^(أولاً|ثانياً|ثالثاً|رابعاً|خامساً|سادساً|سابعاً|ثامناً|تاسعاً|عاشراً):',
        ]

        for pattern in main_patterns:
            if re.search(pattern, text):
                return 'main'

        # FALLBACK: Check for subsection patterns (numbered like 2.1, 2.2)
        # Must start with digit.digit followed by space
        sub_patterns = [
            r'^\d+\.\d+\s',  # Matches: "2.1 ", "3.2 ", etc.
        ]

        for pattern in sub_patterns:
            if re.search(pattern, text):
                return 'sub'

        return None

    def _get_heading_level(self, para) -> int:
        """Get heading level (1-9) from paragraph style."""
        if para.style and 'Heading' in para.style.name:
            match = re.search(r'Heading (\d+)', para.style.name)
            if match:
                return int(match.group(1))

        # Estimate level based on font size
        if para.runs and para.runs[0].font.size:
            size_pt = para.runs[0].font.size.pt
            if size_pt >= 18:
                return 1
            elif size_pt >= 14:
                return 2
            else:
                return 3

        return 2  # Default

    def _generate_section_id(self, title: str, index: int) -> str:
        """
        Generate a unique ID for a section.

        Uses transliteration of Arabic or simple slugification.
        """
        # Simple approach: use index and first few words
        words = title.strip().split()[:3]
        slug = '_'.join(words).lower()

        # Remove special characters
        slug = re.sub(r'[^\w\s_-]', '', slug)
        slug = re.sub(r'[-\s]+', '_', slug)

        return f"section_{index}_{slug}"[:50]  # Limit length

    def _translate_section_title(self, title_ar: str) -> str:
        """
        Simple translation of common Arabic section titles to English.

        For unknown titles, returns transliteration or original.
        """
        translations = {
            'الملخص التنفيذي': 'Executive Summary',
            'المقدمة': 'Introduction',
            'الخلاصة': 'Conclusion',
            'التوصيات': 'Recommendations',
            'حالات الاستخدام': 'Use Cases',
            'خارطة الطريق': 'Roadmap',
        }

        # Check for exact matches
        title_clean = title_ar.strip()
        if title_clean in translations:
            return translations[title_clean]

        # Check for partial matches
        for ar_key, en_val in translations.items():
            if ar_key in title_clean:
                return en_val

        # Check for "Analysis" pattern
        if 'تحليل' in title_clean or 'التحليل' in title_clean:
            # Extract number if present
            match = re.search(r'(الأول|الثاني|الثالث|الرابع|الخامس|\d+)', title_clean)
            if match:
                num_map = {
                    'الأول': 'First',
                    'الثاني': 'Second',
                    'الثالث': 'Third',
                    'الرابع': 'Fourth',
                    'الخامس': 'Fifth'
                }
                num_text = num_map.get(match.group(1), match.group(1))
                return f"Analysis {num_text}"
            return "Analysis"

        # Return original if no translation found
        return title_ar

    def _find_table_position(self, table, table_idx: int) -> int:
        """
        Find the actual position of a table in the document.

        Returns the paragraph index before the table by tracking the actual
        order of elements in the document body.
        """
        # Track all elements (paragraphs and tables) in their actual order
        para_count = 0
        table_count = 0

        for element in self.doc.element.body:
            # Check if this is a table element
            if element.tag.endswith('tbl'):
                # Count which table this is
                if table_count == table_idx:
                    # Found our table - return current paragraph count
                    return para_count
                table_count += 1
            # Check if this is a paragraph
            elif element.tag.endswith('p'):
                para_count += 1

        # If table not found, return approximate position
        total_paras = len(self.doc.paragraphs)
        return total_paras // 2

    def _extract_table_data(self, table) -> Dict[str, Any]:
        """
        Extract complete table data with columns and rows.

        Returns:
            Dictionary with columns, rows, and metadata
        """
        if len(table.rows) == 0:
            return {'columns': [], 'rows': []}

        # Get column headers from first row
        columns = [cell.text.strip() for cell in table.rows[0].cells]

        # Get data rows
        rows = []
        for row in table.rows[1:]:
            row_data = {}
            cells = [cell.text.strip() for cell in row.cells]

            for col_idx, col_name in enumerate(columns):
                if col_idx < len(cells):
                    row_data[col_name] = cells[col_idx]
                else:
                    row_data[col_name] = ''

            rows.append(row_data)

        return {
            'columns': columns,
            'rows': rows,
            'row_count': len(rows),
            'col_count': len(columns)
        }

    def _extract_metadata(self) -> Dict[str, Any]:
        """Extract document metadata."""
        core_props = self.doc.core_properties

        return {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else '',
            'total_paragraphs': len(self.doc.paragraphs),
            'total_tables': len(self.doc.tables)
        }
